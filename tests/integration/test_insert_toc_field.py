# tests/integration/test_insert_toc_field.py
from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document

from docs.infrastructure.docx.python_docx_assembly_adapter import insert_toc_field

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _docx_with_toc_placeholder(tmp_path: Path) -> Path:
    document = Document()
    document.add_paragraph("[[TOC]]")
    path = tmp_path / "fixture.docx"
    document.save(path)
    return path


def test_insert_toc_field_replaces_placeholder_with_toc_field(tmp_path):
    path = _docx_with_toc_placeholder(tmp_path)
    result = insert_toc_field(path)
    assert result is True

    reopened = Document(str(path))
    target = next(p for p in reopened.paragraphs if "TOC" in p._p.xml or "actualizara" in p.text)
    xml = target._p.xml
    assert 'w:fldCharType="begin"' in xml
    assert 'TOC \\o "1-3" \\h \\z \\u' in xml
    assert 'w:fldCharType="separate"' in xml
    assert 'w:fldCharType="end"' in xml


def test_insert_toc_field_sets_update_fields_on_open(tmp_path):
    path = _docx_with_toc_placeholder(tmp_path)
    insert_toc_field(path)

    with zipfile.ZipFile(path) as archive:
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
    assert 'w:val="true"' in settings_xml
    assert "updateFields" in settings_xml


def test_insert_toc_field_returns_false_and_leaves_file_untouched_when_placeholder_missing(tmp_path):
    document = Document()
    document.add_paragraph("No placeholder here.")
    path = tmp_path / "no_toc.docx"
    document.save(path)
    before = path.read_bytes()

    result = insert_toc_field(path)
    assert result is False
    assert path.read_bytes() == before


def test_insert_toc_field_honors_custom_levels_argument(tmp_path):
    path = _docx_with_toc_placeholder(tmp_path)
    insert_toc_field(path, levels="1-2")
    reopened = Document(str(path))
    xml = "".join(p._p.xml for p in reopened.paragraphs)
    assert 'TOC \\o "1-2"' in xml


# --- a placeholder must never survive assembly --------------------------------


def test_a_second_toc_placeholder_is_removed_and_warned_about(tmp_path, capsys):
    # Found by running a real document end to end. `documento-generico`
    # declares BOTH a `{"type": "toc"}` structure part AND an `indice`
    # section whose contract sets `toc: true`, so two `[[TOC]]` paragraphs
    # reach the assembled file. `insert_toc_field` replaced the first and
    # `break`-ed, and the delivered .docx carried a literal `[[TOC]]` under
    # its ÍNDICE heading.
    #
    # The other two builtin templates declare one or the other, never both --
    # so this is a template mistake, but the harness must not emit its own
    # marker as visible text no matter what a template declares.
    document = Document()
    document.add_paragraph("[[TOC]]")
    document.add_heading("ÍNDICE", level=1)
    document.add_paragraph("[[TOC]]")
    path = tmp_path / "dos-toc.docx"
    document.save(path)

    assert insert_toc_field(path) is True

    result = Document(str(path))
    bodies = [p.text.strip() for p in result.paragraphs]
    assert "[[TOC]]" not in bodies, bodies
    assert result.element.body.xml.count("TOC \o") == 1, "un solo campo TOC, no dos"
    assert "WARN" in capsys.readouterr().err


def test_a_single_placeholder_still_becomes_the_field_without_warning(tmp_path, capsys):
    document = Document()
    document.add_paragraph("[[TOC]]")
    path = tmp_path / "un-toc.docx"
    document.save(path)

    assert insert_toc_field(path) is True

    assert "WARN" not in capsys.readouterr().err
