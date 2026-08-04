# tests/integration/test_docx_assembly_service.py
from __future__ import annotations

import json
import shutil
import struct
import subprocess
import zipfile
import zlib
from pathlib import Path

import pytest
from docx import Document

from docs.application.asset import AssetService
from docs.application.docx_assembly import DocxRendererAdapter
from docs.domain.ports.document_renderer_port import DocumentRendererPort
from docs.domain.workspace import Workspace
from docs.infrastructure.docx.python_docx_assembly_adapter import PythonDocxAssemblyAdapter
from docs.infrastructure.docx.tool_resolver_adapter import SystemToolResolverAdapter
from docs.infrastructure.persistence.filesystem_asset_repository import FilesystemAssetRepository


def _png_chunk(tag: bytes, data: bytes) -> bytes:
    return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", zlib.crc32(tag + data))


def _solid_png(width: int, height: int) -> bytes:
    """A minimal, genuinely-parseable RGB PNG at the given size (same
    struct+zlib construction as `test_ingest_assets_figures.py::_solid_png`)
    -- pandoc must be able to embed a real image for the media-entry
    assertions below."""
    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)  # 8-bit RGB
    raw = b"".join(b"\x00" + bytes([180, 180, 180] * width) for _ in range(height))
    idat = zlib.compress(raw)
    return b"\x89PNG\r\n\x1a\n" + _png_chunk(b"IHDR", ihdr) + _png_chunk(b"IDAT", idat) + _png_chunk(b"IEND", b"")


def _write_json(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload), encoding="utf-8")


def _figure_catalog_row(
    catalog_id: str, *, width_px: int | None = 150, height_px: int | None = 100, name: str | None = None
) -> dict:
    return {
        "id": catalog_id,
        "sha256": "0" * 64,
        "width_px": width_px,
        "height_px": height_px,
        "origin_relative_path": f"assets/figures/{name or catalog_id}.png",
        "caption": "Organigrama del equipo",
        "source_role": "evidence",
        "origin_kind": "standalone",
    }


def _pandoc_styled_docx(tmp_path: Path, text: str, name: str) -> Path:
    # A real-world template/cover .docx is a Word document with named styles
    # that match pandoc's docx output (e.g. "First Paragraph", "Body Text") —
    # not a blank python-docx Document(). Generating it via pandoc itself
    # keeps these tests on the realistic path: pandoc always stamps a
    # document's first paragraph with "First Paragraph", a style Word's blank
    # template doesn't define. Mapping/importing missing styles onto an
    # arbitrary template is Slice 11b's `safe_style_name` (stubbed as a no-op
    # in this slice), so a blank-template happy path is not achievable here.
    seed = tmp_path / f"_seed_{name}.md"
    seed.write_text(f"{text}\n", encoding="utf-8")
    target = tmp_path / name
    subprocess.run([shutil.which("pandoc"), str(seed), "-o", str(target)], check=True)
    return target


@pytest.fixture
def workspace(tmp_path: Path) -> Workspace:
    return Workspace(documents_dir=tmp_path / "documents", templates_dir=tmp_path / "templates")


@pytest.fixture
def asset_service(workspace: Workspace) -> AssetService:
    return AssetService(FilesystemAssetRepository(), workspace)


@pytest.fixture
def service(asset_service: AssetService) -> DocxRendererAdapter:
    return DocxRendererAdapter(PythonDocxAssemblyAdapter(), asset_service, SystemToolResolverAdapter())


# --- DocumentRendererPort contract ----------------------------------------------


def test_docx_renderer_adapter_declares_docx_output_format(service):
    assert service.output_format == "docx"


def test_docx_renderer_adapter_satisfies_document_renderer_port(service: DocumentRendererPort):
    assert service.output_format == "docx"
    assert service.stage_plan() == [
        ("build-docx", True),
        ("format-audit-docx", True),
        ("qa-docx", True),
    ]


def test_docx_renderer_adapter_resolves_via_registry_by_format(asset_service):
    from docs.cli._shared import resolve_renderer

    adapter = DocxRendererAdapter(PythonDocxAssemblyAdapter(), asset_service, SystemToolResolverAdapter())
    registry = {"docx": adapter}
    resolved = resolve_renderer(registry, "docx")
    assert resolved is adapter


def test_resolve_renderer_raises_clear_error_on_unregistered_format(asset_service):
    from docs.cli._shared import resolve_renderer

    adapter = DocxRendererAdapter(PythonDocxAssemblyAdapter(), asset_service, SystemToolResolverAdapter())
    registry = {"docx": adapter}
    with pytest.raises(ValueError, match="pdf"):
        resolve_renderer(registry, "pdf")


# --- _resolve_cover_asset_path ------------------------------------------------


def test_resolve_cover_asset_path_returns_none_when_no_cover_from_asset_part(service):
    parts = [{"type": "cover_from_template"}, {"type": "sections"}]
    assert service._resolve_cover_asset_path("doc-1", parts) is None


def test_resolve_cover_asset_path_ignores_parts_after_sections(service):
    parts = [{"type": "sections"}, {"type": "cover_from_asset", "asset": "cover"}]
    assert service._resolve_cover_asset_path("doc-1", parts) is None


def test_resolve_cover_asset_path_uses_asset_service_with_given_doc_id(workspace, asset_service, service):
    cover_dir = workspace.assets_dir("doc-1")
    cover_dir.mkdir(parents=True)
    Document().save(cover_dir / "cover.docx")

    parts = [{"type": "cover_from_asset", "asset": "cover"}, {"type": "sections"}]
    result = service._resolve_cover_asset_path("doc-1", parts)
    assert result == asset_service.asset_path("doc-1", "cover")


def test_resolve_cover_asset_path_defaults_asset_name_to_cover(workspace, service):
    cover_dir = workspace.assets_dir("doc-1")
    cover_dir.mkdir(parents=True)
    Document().save(cover_dir / "cover.docx")

    parts = [{"type": "cover_from_asset"}, {"type": "sections"}]
    assert service._resolve_cover_asset_path("doc-1", parts) == cover_dir / "cover.docx"


def test_resolve_cover_asset_path_isolates_by_doc_id(workspace, service):
    doc1_dir = workspace.assets_dir("doc-1")
    doc1_dir.mkdir(parents=True)
    Document().save(doc1_dir / "cover.docx")
    doc2_dir = workspace.assets_dir("doc-2")
    doc2_dir.mkdir(parents=True)
    Document().save(doc2_dir / "cover.docx")

    parts = [{"type": "cover_from_asset", "asset": "cover"}, {"type": "sections"}]
    result_1 = service._resolve_cover_asset_path("doc-1", parts)
    result_2 = service._resolve_cover_asset_path("doc-2", parts)
    assert result_1 == doc1_dir / "cover.docx"
    assert result_2 == doc2_dir / "cover.docx"
    assert result_1 != result_2


# --- _resolve_embed_paths ------------------------------------------------------


def test_resolve_embed_paths_resolves_front_assets(workspace, service):
    assets_dir = workspace.assets_dir("doc-1")
    assets_dir.mkdir(parents=True)
    Document().save(assets_dir / "front.docx")

    parts = [{"type": "embed_docx", "asset": "front"}, {"type": "sections"}]
    result = service._resolve_embed_paths("doc-1", parts, "front")
    assert result == [assets_dir / "front.docx"]


def test_resolve_embed_paths_resolves_back_assets(workspace, service):
    assets_dir = workspace.assets_dir("doc-1")
    assets_dir.mkdir(parents=True)
    Document().save(assets_dir / "back.docx")

    parts = [{"type": "sections"}, {"type": "embed_docx", "asset": "back"}]
    result = service._resolve_embed_paths("doc-1", parts, "back")
    assert result == [assets_dir / "back.docx"]


def test_resolve_embed_paths_ignores_non_embed_parts(service):
    parts = [{"type": "cover_from_asset", "asset": "cover"}, {"type": "sections"}]
    assert service._resolve_embed_paths("doc-1", parts, "front") == []


def test_resolve_embed_paths_raises_when_asset_missing(service):
    parts = [{"type": "embed_docx", "asset": "missing"}, {"type": "sections"}]
    with pytest.raises(FileNotFoundError):
        service._resolve_embed_paths("doc-1", parts, "front")


def test_resolve_embed_paths_isolates_by_doc_id(workspace, service):
    doc1_assets = workspace.assets_dir("doc-1")
    doc1_assets.mkdir(parents=True)
    Document().save(doc1_assets / "front.docx")

    parts = [{"type": "embed_docx", "asset": "front"}, {"type": "sections"}]
    result = service._resolve_embed_paths("doc-1", parts, "front")
    assert result == [doc1_assets / "front.docx"]
    with pytest.raises(FileNotFoundError):
        service._resolve_embed_paths("doc-2", parts, "front")


# --- assemble ------------------------------------------------------------------


def test_assemble_resolves_cover_asset_path_via_asset_service(tmp_path, workspace, service):
    cover_dir = workspace.assets_dir("doc-1")
    cover_dir.mkdir(parents=True)
    cover = Document()
    cover.add_paragraph("COVER MARKER")
    cover.save(cover_dir / "cover.docx")

    config = {"structure": [{"type": "cover_from_asset", "asset": "cover"}, {"type": "sections"}]}
    body = tmp_path / "body.docx"
    Document().save(body)
    output = tmp_path / "out.docx"

    service.assemble("doc-1", config, body, output)

    assert output.exists()
    result = Document(str(output))
    assert any("COVER MARKER" in p.text for p in result.paragraphs)


def test_assemble_raises_when_embed_asset_missing(tmp_path, service):
    config = {"structure": [{"type": "embed_docx", "asset": "missing"}, {"type": "sections"}]}
    body = tmp_path / "body.docx"
    Document().save(body)
    output = tmp_path / "out.docx"

    with pytest.raises(FileNotFoundError):
        service.assemble("doc-1", config, body, output)


def test_assemble_resolves_and_passes_embed_paths_to_port(tmp_path, workspace, service):
    # docxcompose is now a declared, installed dependency (PR1 quick-debt fix);
    # embedding the front asset must succeed end-to-end, proving doc_id-threaded
    # asset resolution reaches real docxcompose composition without error.
    assets_dir = workspace.assets_dir("doc-1")
    assets_dir.mkdir(parents=True)
    Document().save(assets_dir / "front.docx")

    config = {"structure": [{"type": "embed_docx", "asset": "front"}, {"type": "sections"}]}
    body = tmp_path / "body.docx"
    Document().save(body)
    output = tmp_path / "out.docx"

    service.assemble("doc-1", config, body, output)

    assert output.exists()
    Document(str(output))  # must open without raising


# --- assemble: [[pagebreak]] marker (forced Word page break) -------------------


def test_assemble_renders_sole_pagebreak_marker_as_word_page_break(tmp_path, service):
    body = tmp_path / "body.docx"
    doc = Document()
    doc.add_paragraph("A")
    doc.add_paragraph("[[pagebreak]]")
    doc.add_paragraph("B")
    doc.save(body)
    output = tmp_path / "out.docx"

    service.assemble("doc-1", {"structure": [{"type": "sections"}]}, body, output)

    result = Document(str(output))
    assert not any(p.text.strip() == "[[pagebreak]]" for p in result.paragraphs)
    assert any('w:type="page"' in p._p.xml for p in result.paragraphs)


def test_assemble_leaves_inline_pagebreak_marker_as_literal_text(tmp_path, service):
    # Only a paragraph whose ENTIRE trimmed content is `[[pagebreak]]` triggers
    # a break -- a marker mixed with other text is left untouched, same rule
    # as every other `[[...]]` marker family (e.g. `[[TOC]]`).
    body = tmp_path / "body.docx"
    doc = Document()
    doc.add_paragraph("Antes [[pagebreak]] despues")
    doc.save(body)
    output = tmp_path / "out.docx"

    service.assemble("doc-1", {"structure": [{"type": "sections"}]}, body, output)

    result = Document(str(output))
    assert any("Antes [[pagebreak]] despues" in p.text for p in result.paragraphs)


def test_assemble_pagebreak_before_auto_breaking_heading_emits_single_break(tmp_path, service):
    # Regression (Codex-sub review): a `[[pagebreak]]` marker placed directly
    # before a Heading 1 that ALREADY triggers the per-heading auto-break must
    # yield exactly ONE page break, not two stacked breaks (a blank page). This
    # is the intended usage of the marker (forcing a heading onto a fresh page).
    body = tmp_path / "body.docx"
    doc = Document()
    doc.add_heading("Capitulo I", level=1)  # first H1: no auto-break, marks seen
    doc.add_paragraph("texto")
    doc.add_paragraph("[[pagebreak]]")
    doc.add_heading("Capitulo II", level=1)  # H1 after seen -> auto-break fires
    doc.save(body)
    output = tmp_path / "out.docx"

    service.assemble("doc-1", {"structure": [{"type": "sections"}]}, body, output)

    result = Document(str(output))
    breaks = sum(p._p.xml.count('w:type="page"') for p in result.paragraphs)
    assert breaks == 1


def test_assemble_with_pagebreak_marker_is_byte_identical_across_rebuilds(tmp_path, service):
    body = tmp_path / "body.docx"
    doc = Document()
    doc.add_paragraph("A")
    doc.add_paragraph("[[pagebreak]]")
    doc.add_paragraph("B")
    doc.save(body)

    output_1 = tmp_path / "out1.docx"
    output_2 = tmp_path / "out2.docx"
    service.assemble("doc-1", {"structure": [{"type": "sections"}]}, body, output_1)
    service.assemble("doc-1", {"structure": [{"type": "sections"}]}, body, output_2)

    assert output_1.read_bytes() == output_2.read_bytes()


def test_assemble_strips_filesystem_path_from_picture_description(tmp_path, service):
    # Pandoc leaks the image's absolute source path into the picture's
    # non-visual description (`pic:cNvPr@descr`). Left as-is it makes the
    # assembled docx depend on the build machine's paths (breaking byte
    # identity across machines/workspaces) and leaks local directories into
    # the shipped document. The assembly must reduce that description to its
    # bare basename, which is workspace-independent and stable.
    from docx.oxml.ns import qn
    from PIL import Image

    png = tmp_path / "visual-abcd1234.png"
    Image.new("RGB", (12, 12), color=(1, 2, 3)).save(png)
    body = tmp_path / "body.docx"
    doc = Document()
    run = doc.add_paragraph().add_run()
    run.add_picture(str(png))
    cnvpr = next(doc.element.body.iter(qn("pic:cNvPr")))
    cnvpr.set("descr", r"C:\Users\someone\AppData\Local\Temp\r1\assets\figures\visual-abcd1234.png")
    doc.save(body)
    output = tmp_path / "out.docx"

    service.assemble("doc-1", {"structure": [{"type": "sections"}]}, body, output)

    result = Document(str(output))
    descrs = [el.get("descr") for el in result.element.body.iter(qn("pic:cNvPr"))]
    assert descrs == ["visual-abcd1234.png"]
    assert not any("\\" in (d or "") or "/" in (d or "") for d in descrs)


def test_assemble_centers_figures_and_captions_and_keeps_them_together(tmp_path, service):
    # Academic layout: an image and its "Figura N."/"Tabla N." caption must be
    # centered (not left-aligned) and the image paragraph must keep_with_next so
    # the caption never orphans onto the next page. Normal body text is left
    # untouched.
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from PIL import Image

    png = tmp_path / "f.png"
    Image.new("RGB", (24, 24), color=(4, 5, 6)).save(png)
    body = tmp_path / "body.docx"
    doc = Document()
    doc.add_paragraph().add_run().add_picture(str(png))
    doc.add_paragraph("Figura 1. Una figura de prueba.")
    doc.add_paragraph("Un parrafo normal de cuerpo que no debe centrarse.")
    doc.save(body)
    output = tmp_path / "out.docx"

    service.assemble("doc-1", {"structure": [{"type": "sections"}]}, body, output)

    result = Document(str(output))
    img_para = next(p for p in result.paragraphs if p._p.find(".//" + qn("w:drawing")) is not None)
    cap_para = next(p for p in result.paragraphs if (p.text or "").strip().startswith("Figura 1."))
    body_para = next(p for p in result.paragraphs if "parrafo normal" in (p.text or ""))
    assert img_para.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert img_para.paragraph_format.keep_with_next is True
    assert cap_para.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert body_para.alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_assemble_tables_repeat_header_keep_rows_and_bind_caption(tmp_path, service):
    # Table layout: the "Tabla N." caption keeps with its table (never orphans
    # at a page bottom), the header row repeats on every page a table spans, and
    # no row splits across a page boundary.
    from docx.oxml.ns import qn

    body = tmp_path / "body.docx"
    doc = Document()
    doc.add_paragraph("Tabla 1. Una tabla de prueba.")
    table = doc.add_table(rows=3, cols=2)
    for r in range(3):
        for c in range(2):
            table.cell(r, c).text = f"c{r}{c}"
    doc.save(body)
    output = tmp_path / "out.docx"

    service.assemble("doc-1", {"structure": [{"type": "sections"}]}, body, output)

    result = Document(str(output))
    cap = next(p for p in result.paragraphs if (p.text or "").strip().startswith("Tabla 1."))
    assert cap.paragraph_format.keep_with_next is True
    out_table = result.tables[0]
    rows = out_table.rows
    # header (row 0) repeats on page breaks
    assert rows[0]._tr.find(qn("w:trPr")).find(qn("w:tblHeader")) is not None
    # every row is kept intact across page boundaries
    for row in rows:
        assert row._tr.find(qn("w:trPr")).find(qn("w:cantSplit")) is not None


# --- _strip_frontmatter_to_temp -------------------------------------------------


def test_strip_frontmatter_to_temp_removes_frontmatter_block(tmp_path, service):
    section = tmp_path / "001-resumen.md"
    section.write_text('---\n{"title": "Resumen"}\n---\n# Resumen\n\nCuerpo.\n', encoding="utf-8")

    stripped = service._strip_frontmatter_to_temp([section])

    assert len(stripped) == 1
    assert stripped[0] != section
    assert stripped[0].read_text(encoding="utf-8") == "# Resumen\n\nCuerpo.\n"


def test_strip_frontmatter_to_temp_preserves_content_without_frontmatter(tmp_path, service):
    section = tmp_path / "001-resumen.md"
    section.write_text("# Resumen\n\nCuerpo.\n", encoding="utf-8")

    stripped = service._strip_frontmatter_to_temp([section])
    assert stripped[0].read_text(encoding="utf-8") == "# Resumen\n\nCuerpo.\n"


def test_strip_frontmatter_to_temp_handles_multiple_sections_in_order(tmp_path, service):
    first = tmp_path / "001-resumen.md"
    first.write_text('---\n{"title": "Resumen"}\n---\nUno.\n', encoding="utf-8")
    second = tmp_path / "002-intro.md"
    second.write_text("Dos.\n", encoding="utf-8")

    stripped = service._strip_frontmatter_to_temp([first, second])

    assert [path.read_text(encoding="utf-8") for path in stripped] == ["Uno.\n", "Dos.\n"]


# --- _strip_frontmatter_to_temp: figure/table numbering + cross-ref (item H) ---


def test_strip_frontmatter_to_temp_numbers_figure_markers_in_document_order(tmp_path, service):
    first = tmp_path / "001-resumen.md"
    first.write_text("[[figure:organigrama]] Organigrama del equipo.\n", encoding="utf-8")
    second = tmp_path / "002-anexos.md"
    second.write_text("Consulte [[ref:organigrama]] para más detalle.\n", encoding="utf-8")

    stripped = service._strip_frontmatter_to_temp([first, second])

    assert stripped[0].read_text(encoding="utf-8") == "Figura 1. Organigrama del equipo.\n"
    assert stripped[1].read_text(encoding="utf-8") == "Consulte Ver Figura 1 para más detalle.\n"


def test_strip_frontmatter_to_temp_leaves_hardcoded_captions_untouched(tmp_path, service):
    section = tmp_path / "001-resumen.md"
    section.write_text("Figura 1. Ya numerada a mano.\n", encoding="utf-8")

    stripped = service._strip_frontmatter_to_temp([section])

    assert stripped[0].read_text(encoding="utf-8") == "Figura 1. Ya numerada a mano.\n"


def test_strip_frontmatter_to_temp_warns_on_unresolved_ref_never_silent(tmp_path, service, capsys):
    section = tmp_path / "001-resumen.md"
    section.write_text("Consulte [[ref:no-existe]].\n", encoding="utf-8")

    stripped = service._strip_frontmatter_to_temp([section])

    assert stripped[0].read_text(encoding="utf-8") == "Consulte Ver Figura ?.\n"
    captured = capsys.readouterr()
    assert "WARN" in captured.err
    assert "no-existe" in captured.err


# --- build ----------------------------------------------------------------------


def test_build_raises_when_pandoc_unavailable(tmp_path, monkeypatch, service):
    monkeypatch.setattr("shutil.which", lambda name: None)
    config = {"sections": [], "paths": {"sections_dir": str(tmp_path), "output_draft_dir": str(tmp_path)}}
    with pytest.raises(RuntimeError, match="Pandoc"):
        service.build("doc-1", config)


def test_build_raises_when_no_markdown_sections_exist(tmp_path, service):
    config = {
        "sections": [{"id": "resumen", "order": 1}],
        "paths": {"sections_dir": str(tmp_path / "sections"), "output_draft_dir": str(tmp_path / "draft")},
    }
    (tmp_path / "sections").mkdir()
    with pytest.raises(RuntimeError, match="No hay secciones"):
        service.build("doc-1", config)


# --- config-driven output names (PR4: move hardcoded doc names to config) ------


def test_build_default_output_names_are_doc_id_derived(tmp_path, service):
    # No config["output"] key at all — the default derives from the doc id,
    # not a hardcoded "tesina" literal (residual estadia-coupling fix).
    config = {
        "sections": [{"id": "resumen", "order": 1}],
        "paths": {"sections_dir": str(tmp_path / "sections"), "output_draft_dir": str(tmp_path / "draft")},
    }
    (tmp_path / "sections").mkdir()
    with pytest.raises(RuntimeError, match="No hay secciones"):
        service.build("doc-1", config)
    assert service._draft_docx_name("doc-1", config) == "doc-1-draft.docx"
    assert service._body_docx_name("doc-1", config) == "doc-1-body.docx"


def test_build_uses_configured_output_names_when_present(service):
    config = {"output": {"draft_name": "custom-draft.docx", "body_name": "custom-body.docx"}}
    assert service._draft_docx_name("doc-1", config) == "custom-draft.docx"
    assert service._body_docx_name("doc-1", config) == "custom-body.docx"


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_build_produces_docx_at_configured_draft_and_body_names(tmp_path, service):
    # Behavior-level proof for the config-driven output names (fresh-context
    # review finding: the private-helper-only test above would not have
    # caught application/pipeline.py's audit/QA stages still hardcoding the
    # default name — this test exercises the real build() end-to-end).
    sections_dir = tmp_path / "sections"
    sections_dir.mkdir()
    draft_dir = tmp_path / "draft"
    (sections_dir / "001-resumen.md").write_text("# Resumen\n\nContenido.\n", encoding="utf-8")
    template = _pandoc_styled_docx(tmp_path, "Plantilla.", "template.docx")

    config = {
        "sections": [{"id": "resumen", "order": 1}],
        "paths": {
            "sections_dir": str(sections_dir),
            "output_draft_dir": str(draft_dir),
            "template_docx": str(template),
        },
        "output": {"draft_name": "custom-draft.docx", "body_name": "custom-body.docx"},
    }

    output = service.build("doc-1", config)

    assert output == draft_dir / "custom-draft.docx"
    assert output.exists()
    assert (draft_dir / "custom-body.docx").exists()
    assert not (draft_dir / "doc-1-draft.docx").exists()
    assert not (draft_dir / "doc-1-body.docx").exists()


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_build_produces_docx_with_default_output_path(tmp_path, service):
    sections_dir = tmp_path / "sections"
    sections_dir.mkdir()
    draft_dir = tmp_path / "draft"
    (sections_dir / "001-resumen.md").write_text(
        "# Resumen\n\nContenido del resumen.\n", encoding="utf-8"
    )
    template = _pandoc_styled_docx(tmp_path, "Plantilla.", "template.docx")

    config = {
        "sections": [{"id": "resumen", "order": 1}],
        "paths": {
            "sections_dir": str(sections_dir),
            "output_draft_dir": str(draft_dir),
            "template_docx": str(template),
        },
    }

    output = service.build("doc-1", config)

    assert output == draft_dir / "doc-1-draft.docx"
    assert output.exists()
    document = Document(str(output))
    assert any("Contenido del resumen" in p.text for p in document.paragraphs)


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_build_writes_to_custom_output_path_when_given(tmp_path, service):
    sections_dir = tmp_path / "sections"
    sections_dir.mkdir()
    draft_dir = tmp_path / "draft"
    (sections_dir / "001-resumen.md").write_text("# Resumen\n\nTexto.\n", encoding="utf-8")
    template = _pandoc_styled_docx(tmp_path, "Plantilla.", "template.docx")

    config = {
        "sections": [{"id": "resumen", "order": 1}],
        "paths": {
            "sections_dir": str(sections_dir),
            "output_draft_dir": str(draft_dir),
            "template_docx": str(template),
        },
    }
    custom_output = tmp_path / "custom" / "final.docx"

    output = service.build("doc-1", config, output=custom_output)

    assert output == custom_output
    assert output.exists()


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_build_skips_sections_with_no_markdown_file_on_disk(tmp_path, service):
    sections_dir = tmp_path / "sections"
    sections_dir.mkdir()
    draft_dir = tmp_path / "draft"
    (sections_dir / "001-resumen.md").write_text("# Resumen\n\nTexto.\n", encoding="utf-8")
    template = _pandoc_styled_docx(tmp_path, "Plantilla.", "template.docx")

    config = {
        "sections": [
            {"id": "resumen", "order": 1},
            {"id": "no-existe", "order": 2},
        ],
        "paths": {
            "sections_dir": str(sections_dir),
            "output_draft_dir": str(draft_dir),
            "template_docx": str(template),
        },
    }

    output = service.build("doc-1", config)
    assert output.exists()


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_build_resolves_cover_asset_via_asset_service_with_doc_id(tmp_path, workspace, service):
    cover_dir = workspace.assets_dir("doc-1")
    cover_dir.mkdir(parents=True)
    _pandoc_styled_docx(cover_dir, "COVER FROM BUILD", "cover.docx")

    sections_dir = tmp_path / "sections"
    sections_dir.mkdir()
    draft_dir = tmp_path / "draft"
    (sections_dir / "001-resumen.md").write_text("# Resumen\n\nTexto.\n", encoding="utf-8")

    config = {
        "structure": [{"type": "cover_from_asset", "asset": "cover"}, {"type": "sections"}],
        "sections": [{"id": "resumen", "order": 1}],
        "paths": {"sections_dir": str(sections_dir), "output_draft_dir": str(draft_dir)},
    }

    output = service.build("doc-1", config)
    document = Document(str(output))
    assert any("COVER FROM BUILD" in p.text for p in document.paragraphs)


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_build_produces_working_toc_field_not_literal_placeholder(tmp_path, service):
    sections_dir = tmp_path / "sections"
    sections_dir.mkdir()
    (sections_dir / "001-resumen.md").write_text("# Resumen\n\nContenido.\n", encoding="utf-8")
    config = {
        "sections": [{"id": "resumen", "order": 1}],
        "paths": {"sections_dir": str(sections_dir), "output_draft_dir": str(tmp_path / "draft")},
        "structure": [{"type": "cover_from_template"}, {"type": "toc"}, {"type": "sections"}],
    }

    output = service.build("tesina-demo", config)
    result = Document(str(output))
    assert not any(p.text.strip() == "[[TOC]]" for p in result.paragraphs)
    assert any('w:fldCharType="begin"' in p._p.xml for p in result.paragraphs)


# --- build: figure/table numbering + cross-ref wired end-to-end (item H) -------


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_build_numbers_figures_and_resolves_refs_across_sections(tmp_path, service):
    sections_dir = tmp_path / "sections"
    sections_dir.mkdir()
    draft_dir = tmp_path / "draft"
    (sections_dir / "001-resumen.md").write_text(
        "# Resumen\n\n[[figure:organigrama]] Organigrama del equipo.\n", encoding="utf-8"
    )
    (sections_dir / "002-anexos.md").write_text(
        "# Anexos\n\nConsulte [[ref:organigrama]] para más detalle.\n", encoding="utf-8"
    )
    template = _pandoc_styled_docx(tmp_path, "Plantilla.", "template.docx")

    config = {
        "sections": [{"id": "resumen", "order": 1}, {"id": "anexos", "order": 2}],
        "paths": {
            "sections_dir": str(sections_dir),
            "output_draft_dir": str(draft_dir),
            "template_docx": str(template),
        },
    }

    output = service.build("doc-1", config)

    document = Document(str(output))
    texts = [p.text for p in document.paragraphs]
    assert any("Figura 1. Organigrama del equipo." in t for t in texts)
    assert any("Consulte Ver Figura 1 para más detalle." in t for t in texts)


# --- build: bound-figure embedding, degradation, determinism (S4, ADR-4/5/6) ---


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_build_embeds_bound_figure_and_leaves_unbound_label_text_only(tmp_path, workspace, service):
    sections_dir = tmp_path / "sections"
    sections_dir.mkdir()
    draft_dir = tmp_path / "draft"
    (sections_dir / "001-resumen.md").write_text(
        "# Resumen\n\n[[figure:organigrama]] Organigrama del equipo.\n\n"
        "[[figure:no-vinculada]] Otra figura sin vincular.\n",
        encoding="utf-8",
    )
    assets_dir = workspace.assets_dir("doc-1")
    figures_dir = assets_dir / "figures"
    figures_dir.mkdir(parents=True)
    (figures_dir / "fig-abc12345.png").write_bytes(_solid_png(150, 100))
    _write_json(sections_dir / "figure-catalog.json", {"figures": [_figure_catalog_row("fig-abc12345")]})
    _write_json(
        sections_dir / "figure-bindings.json",
        {"schema": 1, "bindings": {"organigrama": "fig-abc12345"}},
    )
    template = _pandoc_styled_docx(tmp_path, "Plantilla.", "template.docx")

    config = {
        "sections": [{"id": "resumen", "order": 1}],
        "paths": {
            "sections_dir": str(sections_dir),
            "output_draft_dir": str(draft_dir),
            "template_docx": str(template),
            "assets_dir": str(assets_dir),
        },
    }

    output = service.build("doc-1", config)

    with zipfile.ZipFile(output) as archive:
        media_entries = [name for name in archive.namelist() if name.startswith("word/media/")]
    # Exactly one image embedded -- the BOUND figure only, the unbound label
    # never reaches pandoc as an image reference.
    assert len(media_entries) == 1

    document = Document(str(output))
    texts = [p.text for p in document.paragraphs]
    assert any("Organigrama del equipo" in t for t in texts)
    assert any("Figura 2. Otra figura sin vincular." in t for t in texts)


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_build_degrades_gracefully_when_bound_image_file_is_missing(tmp_path, workspace, service, capsys):
    sections_dir = tmp_path / "sections"
    sections_dir.mkdir()
    draft_dir = tmp_path / "draft"
    (sections_dir / "001-resumen.md").write_text(
        "# Resumen\n\n[[figure:organigrama]] Organigrama del equipo.\n", encoding="utf-8"
    )
    assets_dir = workspace.assets_dir("doc-1")
    # No file written under assets_dir/figures/ -- the bound image is missing.
    _write_json(sections_dir / "figure-catalog.json", {"figures": [_figure_catalog_row("fig-abc12345")]})
    _write_json(
        sections_dir / "figure-bindings.json",
        {"schema": 1, "bindings": {"organigrama": "fig-abc12345"}},
    )
    template = _pandoc_styled_docx(tmp_path, "Plantilla.", "template.docx")

    config = {
        "sections": [{"id": "resumen", "order": 1}],
        "paths": {
            "sections_dir": str(sections_dir),
            "output_draft_dir": str(draft_dir),
            "template_docx": str(template),
            "assets_dir": str(assets_dir),
        },
    }

    output = service.build("doc-1", config)  # must not raise

    with zipfile.ZipFile(output) as archive:
        media_entries = [name for name in archive.namelist() if name.startswith("word/media/")]
    assert media_entries == []
    document = Document(str(output))
    assert any("Figura 1. Organigrama del equipo." in p.text for p in document.paragraphs)
    captured = capsys.readouterr()
    assert "WARN" in captured.err
    assert "organigrama" in captured.err


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_build_degrades_gracefully_for_corrupt_but_present_bound_image(tmp_path, workspace, service, capsys):
    # "Corrupt-but-present" (tasks.md S4 4.6) maps to the SAME signal the
    # ingest layer already uses for a corrupted image (`_read_image_dimensions`
    # graceful degradation, ingest.py): the raw file DID get copied to
    # `assets_dir/figures/` (present), but its catalog row carries null
    # `width_px`/`height_px` because ingest's dimension read failed on it.
    # ADR-6 explicitly reuses that catalog signal instead of re-opening the
    # file in the renderer (no new port dependency here), so this is the
    # resolver's existing null-dims guard exercised against a present file.
    sections_dir = tmp_path / "sections"
    sections_dir.mkdir()
    draft_dir = tmp_path / "draft"
    (sections_dir / "001-resumen.md").write_text(
        "# Resumen\n\n[[figure:organigrama]] Organigrama del equipo.\n", encoding="utf-8"
    )
    assets_dir = workspace.assets_dir("doc-1")
    figures_dir = assets_dir / "figures"
    figures_dir.mkdir(parents=True)
    (figures_dir / "fig-abc12345.png").write_bytes(b"not-a-real-image-just-garbage-bytes")
    _write_json(
        sections_dir / "figure-catalog.json",
        {"figures": [_figure_catalog_row("fig-abc12345", width_px=None, height_px=None)]},
    )
    _write_json(
        sections_dir / "figure-bindings.json",
        {"schema": 1, "bindings": {"organigrama": "fig-abc12345"}},
    )
    template = _pandoc_styled_docx(tmp_path, "Plantilla.", "template.docx")

    config = {
        "sections": [{"id": "resumen", "order": 1}],
        "paths": {
            "sections_dir": str(sections_dir),
            "output_draft_dir": str(draft_dir),
            "template_docx": str(template),
            "assets_dir": str(assets_dir),
        },
    }

    output = service.build("doc-1", config)  # must not raise

    with zipfile.ZipFile(output) as archive:
        media_entries = [name for name in archive.namelist() if name.startswith("word/media/")]
    assert media_entries == []
    document = Document(str(output))
    assert any("Figura 1. Organigrama del equipo." in p.text for p in document.paragraphs)
    captured = capsys.readouterr()
    assert "WARN" in captured.err
    assert "sin dimensiones" in captured.err


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_build_embeds_healthy_figures_when_one_bound_figure_among_several_is_degraded(
    tmp_path, workspace, service, capsys
):
    sections_dir = tmp_path / "sections"
    sections_dir.mkdir()
    draft_dir = tmp_path / "draft"
    (sections_dir / "001-resumen.md").write_text(
        "# Resumen\n\n[[figure:sana-uno]] Primera figura sana.\n\n"
        "[[figure:degradada]] Figura degradada.\n\n"
        "[[figure:sana-dos]] Segunda figura sana.\n",
        encoding="utf-8",
    )
    assets_dir = workspace.assets_dir("doc-1")
    figures_dir = assets_dir / "figures"
    figures_dir.mkdir(parents=True)
    (figures_dir / "fig-sana0001.png").write_bytes(_solid_png(150, 100))
    (figures_dir / "fig-sana0002.png").write_bytes(_solid_png(160, 110))
    # `fig-degradada.png` intentionally never written -- missing bound image.
    _write_json(
        sections_dir / "figure-catalog.json",
        {
            "figures": [
                _figure_catalog_row("fig-sana0001"),
                _figure_catalog_row("fig-degradada"),
                _figure_catalog_row("fig-sana0002"),
            ]
        },
    )
    _write_json(
        sections_dir / "figure-bindings.json",
        {
            "schema": 1,
            "bindings": {
                "sana-uno": "fig-sana0001",
                "degradada": "fig-degradada",
                "sana-dos": "fig-sana0002",
            },
        },
    )
    template = _pandoc_styled_docx(tmp_path, "Plantilla.", "template.docx")

    config = {
        "sections": [{"id": "resumen", "order": 1}],
        "paths": {
            "sections_dir": str(sections_dir),
            "output_draft_dir": str(draft_dir),
            "template_docx": str(template),
            "assets_dir": str(assets_dir),
        },
    }

    output = service.build("doc-1", config)  # must not raise despite the degraded figure

    with zipfile.ZipFile(output) as archive:
        media_entries = [name for name in archive.namelist() if name.startswith("word/media/")]
    # Both healthy figures embed; the degraded one contributes no media entry.
    assert len(media_entries) == 2
    document = Document(str(output))
    texts = [p.text for p in document.paragraphs]
    assert any("Figura 2. Figura degradada." in t for t in texts)
    captured = capsys.readouterr()
    assert "WARN" in captured.err
    assert "degradada" in captured.err
