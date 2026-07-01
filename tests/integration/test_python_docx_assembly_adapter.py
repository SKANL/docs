# tests/integration/test_python_docx_assembly_adapter.py
from __future__ import annotations

import shutil
import subprocess
from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from docs.infrastructure.docx import python_docx_assembly_adapter
from docs.infrastructure.docx.python_docx_assembly_adapter import (
    PythonDocxAssemblyAdapter,
    configure_numbered_body_section,
    configure_roman_preliminary_section,
)


# --- render_pandoc -----------------------------------------------------------


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_render_pandoc_converts_markdown_to_docx(tmp_path):
    markdown = tmp_path / "section.md"
    markdown.write_text("# Título\n\nCuerpo del texto.\n", encoding="utf-8")
    output = tmp_path / "body.docx"
    PythonDocxAssemblyAdapter().render_pandoc(shutil.which("pandoc"), [markdown], output)
    assert output.exists()
    document = Document(str(output))
    assert any("Cuerpo del texto" in p.text for p in document.paragraphs)


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_render_pandoc_raises_on_pandoc_failure(tmp_path):
    missing_input = tmp_path / "does-not-exist.md"
    output = tmp_path / "body.docx"
    with pytest.raises(subprocess.CalledProcessError):
        PythonDocxAssemblyAdapter().render_pandoc(shutil.which("pandoc"), [missing_input], output)


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_assemble_does_not_crash_on_real_pandoc_first_paragraph_style(tmp_path):
    # Regression test for the gap flagged during Slice 11a review: real
    # pandoc output stamps the paragraph right after a heading with style
    # "First Paragraph", which a blank cover Document() does not define.
    # Before this slice's safe_style_name, this raised KeyError via
    # document.add_paragraph(style="First Paragraph").
    markdown = tmp_path / "section.md"
    markdown.write_text("# Titulo\n\nCuerpo del texto de prueba.\n", encoding="utf-8")
    body_docx = tmp_path / "body.docx"
    PythonDocxAssemblyAdapter().render_pandoc(shutil.which("pandoc"), [markdown], body_docx)

    output = tmp_path / "out.docx"
    # Should not raise KeyError.
    PythonDocxAssemblyAdapter().assemble(
        {}, body_docx, output, cover_asset_path=None, embed_front_paths=[], embed_back_paths=[]
    )
    result = Document(str(output))
    target = next(p for p in result.paragraphs if p.text.strip() == "Cuerpo del texto de prueba.")
    assert target.style.name == "No Spacing"


# --- fixtures ------------------------------------------------------------------


def _save_body_docx(tmp_path: Path, name: str = "body.docx") -> Path:
    document = Document()
    document.add_heading("Introduccion", level=1)
    document.add_paragraph("Texto de cuerpo.")
    path = tmp_path / name
    document.save(path)
    return path


def _count_page_breaks(document) -> int:
    return sum(1 for paragraph in document.paragraphs if 'w:type="page"' in paragraph._p.xml)


# --- assemble: cover resolution -------------------------------------------------


def test_assemble_produces_output_with_blank_cover_when_no_template(tmp_path):
    body = _save_body_docx(tmp_path)
    output = tmp_path / "out.docx"
    config: dict = {}
    PythonDocxAssemblyAdapter().assemble(
        config, body, output, cover_asset_path=None, embed_front_paths=[], embed_back_paths=[]
    )
    assert output.exists()
    document = Document(str(output))
    assert any("Introduccion" in p.text for p in document.paragraphs)


def test_assemble_loads_cover_from_template_when_configured(tmp_path):
    template = Document()
    template.add_paragraph("TEMPLATE COVER MARKER")
    template_path = tmp_path / "template.docx"
    template.save(template_path)

    body = _save_body_docx(tmp_path)
    output = tmp_path / "out.docx"
    config = {"paths": {"template_docx": str(template_path)}}
    PythonDocxAssemblyAdapter().assemble(
        config, body, output, cover_asset_path=None, embed_front_paths=[], embed_back_paths=[]
    )
    document = Document(str(output))
    assert any("TEMPLATE COVER MARKER" in p.text for p in document.paragraphs)


def test_assemble_loads_cover_from_asset_when_structure_declares_it(tmp_path):
    cover = Document()
    cover.add_paragraph("COVER ASSET MARKER")
    cover_path = tmp_path / "cover.docx"
    cover.save(cover_path)

    body = _save_body_docx(tmp_path)
    output = tmp_path / "out.docx"
    config = {"structure": [{"type": "cover_from_asset", "asset": "cover"}, {"type": "sections"}]}
    PythonDocxAssemblyAdapter().assemble(
        config, body, output, cover_asset_path=cover_path, embed_front_paths=[], embed_back_paths=[]
    )
    document = Document(str(output))
    assert any("COVER ASSET MARKER" in p.text for p in document.paragraphs)


def test_assemble_falls_back_to_blank_cover_when_asset_path_missing(tmp_path):
    body = _save_body_docx(tmp_path)
    output = tmp_path / "out.docx"
    config = {"structure": [{"type": "cover_from_asset", "asset": "cover"}, {"type": "sections"}]}
    PythonDocxAssemblyAdapter().assemble(
        config, body, output, cover_asset_path=None, embed_front_paths=[], embed_back_paths=[]
    )
    document = Document(str(output))
    assert not any("COVER ASSET MARKER" in p.text for p in document.paragraphs)
    assert any("Introduccion" in p.text for p in document.paragraphs)


# --- assemble: body traversal ---------------------------------------------------


def test_assemble_copies_tables_from_body(tmp_path):
    document = Document()
    document.add_heading("Capitulo", level=1)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "celda-a"
    table.cell(0, 1).text = "celda-b"
    body = tmp_path / "body.docx"
    document.save(body)

    output = tmp_path / "out.docx"
    PythonDocxAssemblyAdapter().assemble({}, body, output, cover_asset_path=None, embed_front_paths=[], embed_back_paths=[])
    result = Document(str(output))
    assert len(result.tables) == 1
    assert result.tables[0].cell(0, 0).text == "celda-a"


def test_assemble_applies_normative_paragraph_format_line_spacing(tmp_path):
    body = _save_body_docx(tmp_path)
    output = tmp_path / "out.docx"
    PythonDocxAssemblyAdapter().assemble({}, body, output, cover_asset_path=None, embed_front_paths=[], embed_back_paths=[])
    document = Document(str(output))
    body_paragraphs = [p for p in document.paragraphs if p.text.strip() == "Texto de cuerpo."]
    assert body_paragraphs[0].paragraph_format.line_spacing == 1.5


def test_assemble_centers_heading_1_paragraphs(tmp_path):
    body = _save_body_docx(tmp_path)
    output = tmp_path / "out.docx"
    PythonDocxAssemblyAdapter().assemble({}, body, output, cover_asset_path=None, embed_front_paths=[], embed_back_paths=[])
    document = Document(str(output))
    headings = [p for p in document.paragraphs if p.text.strip() == "Introduccion" and p.style and p.style.name == "Heading 1"]
    assert headings
    assert headings[0].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_assemble_copies_run_formatting_from_body(tmp_path):
    document = Document()
    document.add_heading("Capitulo", level=1)
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Texto en negrita.")
    run.bold = True
    body = tmp_path / "body.docx"
    document.save(body)

    output = tmp_path / "out.docx"
    PythonDocxAssemblyAdapter().assemble({}, body, output, cover_asset_path=None, embed_front_paths=[], embed_back_paths=[])
    result = Document(str(output))
    target = next(p for p in result.paragraphs if p.text.strip() == "Texto en negrita.")
    assert target.runs[0].bold is True
    assert target.runs[0].font.name == "Times New Roman"
    assert target.runs[0].font.size == Pt(12)


def test_assemble_inserts_page_break_before_second_heading_1(tmp_path):
    document = Document()
    document.add_heading("Primero", level=1)
    document.add_paragraph("Texto uno.")
    document.add_heading("Segundo", level=1)
    document.add_paragraph("Texto dos.")
    body = tmp_path / "body.docx"
    document.save(body)

    output = tmp_path / "out.docx"
    PythonDocxAssemblyAdapter().assemble({}, body, output, cover_asset_path=None, embed_front_paths=[], embed_back_paths=[])
    result = Document(str(output))
    assert _count_page_breaks(result) == 1


# --- assemble: section count (structural only, not numbering format) ------------


def _restart_config() -> dict:
    return {
        "structure": [
            {"type": "cover_from_template"},
            {
                "type": "sections",
                "preliminary_pagination": {},
                "body_restart_section": "cap2",
                "body_pagination": {"format": "decimal", "start": 1},
            },
        ],
        "sections": [{"id": "cap2", "title": "CAPITULO DOS"}],
    }


def test_assemble_adds_extra_section_when_restart_heading_matches(tmp_path):
    document = Document()
    document.add_heading("CAPITULO DOS", level=1)
    document.add_paragraph("Texto de cuerpo.")
    body = tmp_path / "body.docx"
    document.save(body)

    output = tmp_path / "out.docx"
    PythonDocxAssemblyAdapter().assemble(
        _restart_config(), body, output, cover_asset_path=None, embed_front_paths=[], embed_back_paths=[]
    )
    result = Document(str(output))
    assert len(result.sections) == 3


def test_assemble_does_not_add_extra_section_when_restart_heading_does_not_match(tmp_path):
    body = _save_body_docx(tmp_path)
    output = tmp_path / "out.docx"
    PythonDocxAssemblyAdapter().assemble(
        _restart_config(), body, output, cover_asset_path=None, embed_front_paths=[], embed_back_paths=[]
    )
    result = Document(str(output))
    assert len(result.sections) == 2


# --- assemble: embed front/back + docxcompose missing ---------------------------


def test_assemble_raises_runtime_error_when_docxcompose_missing_with_front_asset(tmp_path, monkeypatch):
    body = _save_body_docx(tmp_path)
    front_asset = tmp_path / "front.docx"
    Document().save(front_asset)
    output = tmp_path / "out.docx"
    monkeypatch.setitem(__import__("sys").modules, "docxcompose", None)
    with pytest.raises(RuntimeError, match="docxcompose"):
        PythonDocxAssemblyAdapter().assemble(
            {}, body, output, cover_asset_path=None, embed_front_paths=[front_asset], embed_back_paths=[]
        )


def test_assemble_raises_runtime_error_when_docxcompose_missing_with_back_asset(tmp_path, monkeypatch):
    body = _save_body_docx(tmp_path)
    back_asset = tmp_path / "back.docx"
    Document().save(back_asset)
    output = tmp_path / "out.docx"
    monkeypatch.setitem(__import__("sys").modules, "docxcompose", None)
    with pytest.raises(RuntimeError, match="docxcompose"):
        PythonDocxAssemblyAdapter().assemble(
            {}, body, output, cover_asset_path=None, embed_front_paths=[], embed_back_paths=[back_asset]
        )


# --- assemble: output path handling ----------------------------------------------


def test_assemble_creates_output_parent_directory_when_missing(tmp_path):
    body = _save_body_docx(tmp_path)
    output = tmp_path / "nested" / "dir" / "out.docx"
    PythonDocxAssemblyAdapter().assemble({}, body, output, cover_asset_path=None, embed_front_paths=[], embed_back_paths=[])
    assert output.exists()


# --- assemble: section pagination (Word-correct, Slice 11b) ---------------------


def _config_with_roman_prelim_and_body_restart() -> dict:
    return {
        "structure": [
            {"type": "cover_from_template"},
            {
                "type": "sections",
                "preliminary_pagination": {"start": 2, "format": "lowerRoman"},
                "body_restart_section": "cap2",
                "body_pagination": {"format": "decimal", "start": 1},
            },
        ],
        "sections": [{"id": "cap2", "title": "CAPITULO DOS"}],
    }


def test_assemble_configures_lower_roman_preliminary_section_pagination(tmp_path):
    document = Document()
    document.add_heading("CAPITULO DOS", level=1)
    document.add_paragraph("Texto de cuerpo.")
    body = tmp_path / "body.docx"
    document.save(body)

    output = tmp_path / "out.docx"
    PythonDocxAssemblyAdapter().assemble(
        _config_with_roman_prelim_and_body_restart(), body, output,
        cover_asset_path=None, embed_front_paths=[], embed_back_paths=[],
    )
    result = Document(str(output))
    prelim_section = result.sections[1]
    pg_num_type = prelim_section._sectPr.find(qn("w:pgNumType"))
    assert pg_num_type.get(qn("w:start")) == "2"
    assert pg_num_type.get(qn("w:fmt")) == "lowerRoman"


def test_assemble_configures_decimal_body_section_pagination_restart(tmp_path):
    document = Document()
    document.add_heading("CAPITULO DOS", level=1)
    document.add_paragraph("Texto de cuerpo.")
    body = tmp_path / "body.docx"
    document.save(body)

    output = tmp_path / "out.docx"
    PythonDocxAssemblyAdapter().assemble(
        _config_with_roman_prelim_and_body_restart(), body, output,
        cover_asset_path=None, embed_front_paths=[], embed_back_paths=[],
    )
    result = Document(str(output))
    body_section = result.sections[2]
    pg_num_type = body_section._sectPr.find(qn("w:pgNumType"))
    assert pg_num_type.get(qn("w:start")) == "1"
    assert pg_num_type.get(qn("w:fmt")) == "decimal"
    footer_xml = body_section.footer.paragraphs[-1]._p.xml
    assert "PAGE" in footer_xml


def test_assemble_unnumbered_section_has_no_page_field_in_footer(tmp_path):
    body = _save_body_docx(tmp_path)
    output = tmp_path / "out.docx"
    config = {
        "structure": [
            {"type": "cover_from_template"},
            {"type": "sections", "preliminary_pagination": {}},
        ]
    }
    PythonDocxAssemblyAdapter().assemble(
        config, body, output, cover_asset_path=None, embed_front_paths=[], embed_back_paths=[]
    )
    result = Document(str(output))
    prelim_section = result.sections[1]
    assert "PAGE" not in prelim_section.footer.paragraphs[-1]._p.xml


def test_assemble_clears_header_and_footer_on_configured_sections(tmp_path):
    body = _save_body_docx(tmp_path)
    output = tmp_path / "out.docx"
    PythonDocxAssemblyAdapter().assemble(
        _config_with_roman_prelim_and_body_restart(), body, output,
        cover_asset_path=None, embed_front_paths=[], embed_back_paths=[],
    )
    result = Document(str(output))
    prelim_section = result.sections[1]
    assert prelim_section.header.is_linked_to_previous is False
    assert prelim_section.footer.is_linked_to_previous is False


# --- configure_*_section: direct unit coverage against transposition (Slice 11b fix) ---
#
# Fresh-context review of commit 237b773 found that the two pipeline-level tests
# above (test_assemble_configures_lower_roman_preliminary_section_pagination and
# test_assemble_configures_decimal_body_section_pagination_restart) do NOT catch a
# copy-paste transposition between configure_roman_preliminary_section and
# configure_numbered_body_section, because _build_main_document unconditionally
# re-calls set_section_page_number_start right after each configure_*_section call,
# overwriting whatever the (possibly wrong) configure_* function set. The tests
# below close that gap: they call each configure_*_section function directly on a
# bare Section (bypassing _build_main_document's redundant overwrite entirely), and
# they spy on which configurator fires for which branch inside assemble().


def test_configure_roman_preliminary_section_sets_lower_roman_pagination_directly():
    document = Document()
    section = document.add_section(WD_SECTION_START.NEW_PAGE)
    configure_roman_preliminary_section(section, {}, start=2)
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    assert pg_num_type.get(qn("w:start")) == "2"
    assert pg_num_type.get(qn("w:fmt")) == "lowerRoman"


def test_configure_numbered_body_section_sets_decimal_pagination_directly():
    document = Document()
    section = document.add_section(WD_SECTION_START.NEW_PAGE)
    configure_numbered_body_section(section, {})
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    assert pg_num_type.get(qn("w:start")) == "1"
    assert pg_num_type.get(qn("w:fmt")) == "decimal"


def test_assemble_invokes_configure_functions_in_correct_order_for_correct_branch(tmp_path, monkeypatch):
    # Spies on the two module-level configurator names so we can assert WHICH
    # function fires for WHICH section-construction branch, independent of what
    # that function's pagination side effects produce (those get overwritten by
    # the redundant set_section_page_number_start call right after, see above).
    call_order: list[str] = []
    original_roman = python_docx_assembly_adapter.configure_roman_preliminary_section
    original_numbered = python_docx_assembly_adapter.configure_numbered_body_section

    def roman_spy(section, config, start=2):
        call_order.append("roman_preliminary")
        return original_roman(section, config, start)

    def numbered_spy(section, config):
        call_order.append("numbered_body")
        return original_numbered(section, config)

    monkeypatch.setattr(python_docx_assembly_adapter, "configure_roman_preliminary_section", roman_spy)
    monkeypatch.setattr(python_docx_assembly_adapter, "configure_numbered_body_section", numbered_spy)

    document = Document()
    document.add_heading("CAPITULO DOS", level=1)
    document.add_paragraph("Texto de cuerpo.")
    body = tmp_path / "body.docx"
    document.save(body)

    output = tmp_path / "out.docx"
    PythonDocxAssemblyAdapter().assemble(
        _config_with_roman_prelim_and_body_restart(), body, output,
        cover_asset_path=None, embed_front_paths=[], embed_back_paths=[],
    )

    # The preliminary section is built first and must use the roman configurator;
    # the body-restart section is built afterward and must use the numbered
    # configurator. A transposed call site would flip this order.
    assert call_order == ["roman_preliminary", "numbered_body"]
