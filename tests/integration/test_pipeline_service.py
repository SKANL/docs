# tests/integration/test_pipeline_service.py
from __future__ import annotations

import json
import shutil
from pathlib import Path

import pytest
from docx import Document

from docs.application.collection import CollectionService
from docs.application.context_pack import ContextPackService
from docs.application.doctor import DoctorService
from docs.application.docx_assembly import DocxRendererAdapter
from docs.application.evidence import EvidenceService
from docs.application.format_audit import FormatAuditService
from docs.application.ingest import IngestService
from docs.application.pipeline import PipelineService
from docs.application.qa import QaService
from docs.application.review import ReviewService
from docs.domain.context import TopicStatus
from docs.domain.models.template import ContextSchema, Template, Topic
from docs.domain.workspace import Workspace
from docs.infrastructure.docx.libreoffice_qa_adapter import LibreOfficeQaAdapter
from docs.infrastructure.docx.python_docx_assembly_adapter import PythonDocxAssemblyAdapter
from docs.infrastructure.docx.python_docx_audit_adapter import PythonDocxAuditAdapter
from docs.infrastructure.docx.tool_resolver_adapter import SystemToolResolverAdapter
from docs.infrastructure.ingest.filetype_detector_adapter import FiletypeDetectorAdapter
from docs.infrastructure.ingest.md_normalize_adapter import MdNormalizeAdapter
from docs.infrastructure.ingest.opendataloader_pdf_adapter import OpendataloaderPdfAdapter
from docs.infrastructure.ingest.pandoc_ingest_adapter import PandocIngestAdapter
from docs.infrastructure.persistence.filesystem_asset_repository import FilesystemAssetRepository
from docs.infrastructure.persistence.filesystem_source_repository import FilesystemSourceRepository
from docs.infrastructure.persistence.json_context_repository import JsonContextRepository
from docs.infrastructure.persistence.json_evidence_repository import JsonEvidenceRepository
from docs.infrastructure.persistence.json_section_repository import JsonSectionRepository
from docs.application.asset import AssetService

_HAS_LIBREOFFICE = shutil.which("soffice") is not None or shutil.which("libreoffice") is not None


def _service(tmp_path) -> tuple[PipelineService, Workspace]:
    workspace = Workspace(documents_dir=tmp_path / "documents", templates_dir=tmp_path / "templates")
    evidence_repo = JsonEvidenceRepository()
    section_repo = JsonSectionRepository(workspace)
    source_repo = FilesystemSourceRepository()
    context_repo = JsonContextRepository(workspace)
    asset_service = AssetService(FilesystemAssetRepository(), workspace)
    evidence_service = EvidenceService(evidence_repo)
    review_service = ReviewService(section_repo)
    collection_service = CollectionService(source_repo, evidence_repo)
    context_pack_service = ContextPackService(section_repo, evidence_repo, evidence_service, review_service)
    tool_resolver = SystemToolResolverAdapter()
    docx_assembly_service = DocxRendererAdapter(PythonDocxAssemblyAdapter(), asset_service, tool_resolver)
    format_audit_service = FormatAuditService(PythonDocxAuditAdapter())
    qa_service = QaService(LibreOfficeQaAdapter(), format_audit_service)
    doctor_service = DoctorService(evidence_repo, asset_service, tool_resolver)
    pandoc_ingest_adapter = PandocIngestAdapter(tool_resolver)
    pdf_ingest_adapter = OpendataloaderPdfAdapter(tool_resolver)
    md_ingest_adapter = MdNormalizeAdapter()
    ingest_service = IngestService(
        FiletypeDetectorAdapter(),
        {
            "docx": pandoc_ingest_adapter,
            "odt": pandoc_ingest_adapter,
            "pdf": pdf_ingest_adapter,
            "md": md_ingest_adapter,
            "txt": md_ingest_adapter,
        },
    )
    service = PipelineService(
        doctor_service, evidence_service, evidence_repo, collection_service, source_repo,
        review_service, context_pack_service, context_repo, docx_assembly_service,
        format_audit_service, qa_service, workspace, ingest_service,
    )
    return service, workspace


def test_build_section_renders_scaffold_gathers_six_hashes_and_writes_section_file(tmp_path: Path):
    from docs.domain.models.template import ContextSchema, Field, Section, SectionContract, Topic

    service, workspace = _service(tmp_path)
    topic = Topic(id="alumno", title="Alumno", consumed_by=["introduccion"], fields=[Field(key="nombre", label="Nombre")])
    template = Template(
        type="tesina",
        title="Tesina",
        sections=[Section(id="introduccion", title="Introducción", order=1, required=True)],
        section_contracts={"introduccion": SectionContract(required_content=["alcance"])},
        context_schema=ContextSchema(topics=[topic]),
    )
    service.context_repository.write_topic("doc-1", topic, {"nombre": "Ana"})
    config = {
        "paths": {
            "manual_dir": str(tmp_path / "manual"),
            "extracted_dir": str(tmp_path / "extracted"),
            "rules_manifest": str(tmp_path / "manual-rules.json"),
            "context_dir": str(tmp_path / "context"),
            "prompts_dir": str(tmp_path / "prompts"),
            "source_manifest": str(tmp_path / "source-manifest.json"),
            "code_evidence_manifest": str(tmp_path / "code-evidence-manifest.json"),
        },
        "sections": [{"id": "introduccion"}],
        "section_contracts": {"introduccion": {"required_content": ["alcance"]}},
        "format": {},
        "apa7": {},
        "structure": [],
        "preliminaries": {},
    }

    path = service.build_section("doc-1", template, "introduccion", config)

    assert path.exists()
    raw = path.read_text(encoding="utf-8")
    assert "- Nombre: Ana" in raw
    assert "PENDIENTE: documentar alcance con evidencia del ledger, contexto o fuentes." in raw
    metadata = json.loads(raw.split("---\n")[1])
    assert metadata["section_id"] == "introduccion"
    assert len(metadata["source_hash"]) == 64
    assert len(metadata["prompt_hash"]) == 64
    assert len(metadata["rules_hash"]) == 64
    assert len(metadata["contract_hash"]) == 64
    assert metadata["source_manifest_hash"] == ""  # manifest never built -> manifest_hash("") sentinel
    assert metadata["code_evidence_manifest_hash"] == ""


def test_build_section_only_includes_context_topics_consumed_by_the_target_section(tmp_path: Path):
    from docs.domain.models.template import ContextSchema, Section, SectionContract, Topic

    service, workspace = _service(tmp_path)
    other_topic = Topic(id="otro", title="Otro", consumed_by=["otra-seccion"], multiline=True)
    template = Template(
        type="tesina",
        title="Tesina",
        sections=[Section(id="introduccion", title="Introducción", order=1, required=True)],
        section_contracts={"introduccion": SectionContract()},
        context_schema=ContextSchema(topics=[other_topic]),
    )
    service.context_repository.write_topic("doc-1", other_topic, "Texto no relacionado con introduccion.")
    config = {
        "paths": {
            "manual_dir": str(tmp_path / "manual"),
            "extracted_dir": str(tmp_path / "extracted"),
            "rules_manifest": str(tmp_path / "manual-rules.json"),
            "context_dir": str(tmp_path / "context"),
            "prompts_dir": str(tmp_path / "prompts"),
            "source_manifest": str(tmp_path / "source-manifest.json"),
            "code_evidence_manifest": str(tmp_path / "code-evidence-manifest.json"),
        },
        "sections": [{"id": "introduccion"}],
        "section_contracts": {},
        "format": {},
        "apa7": {},
        "structure": [],
        "preliminaries": {},
    }

    path = service.build_section("doc-1", template, "introduccion", config)

    raw = path.read_text(encoding="utf-8")
    assert "Texto no relacionado" not in raw


def test_build_section_raises_file_not_found_for_unknown_section_id(tmp_path: Path):
    service, _ = _service(tmp_path)
    template = Template(type="tesina", title="Tesina", sections=[])

    with pytest.raises(FileNotFoundError, match="No existe sección: no-existe"):
        service.build_section("doc-1", template, "no-existe", {"paths": {}})


def test_rules_manifest_state_goes_through_evidence_repository_not_direct_stat(tmp_path, monkeypatch):
    # rules_path is never created on disk: a Path.stat()-direct implementation
    # would see it as absent (exists=False, size=0) and could not possibly
    # reproduce the values below. The injected evidence_repository fake
    # reports contradictory values (exists=True, size=999) for that same
    # missing path. If rules_manifest_state() still returns (True, 999), the
    # only way that is possible is that it called through
    # self.evidence_repository.file_exists/file_size rather than touching the
    # filesystem itself.
    service, workspace = _service(tmp_path)
    rules_path = tmp_path / "manual-rules.json"
    assert not rules_path.exists()
    monkeypatch.setattr(service.evidence_repository, "file_exists", lambda path: True)
    monkeypatch.setattr(service.evidence_repository, "file_size", lambda path: 999)
    config = {"paths": {"rules_manifest": str(rules_path)}}

    exists, size = service.rules_manifest_state(config)

    assert exists is True
    assert size == 999


def test_rules_manifest_state_reports_absent_manifest(tmp_path):
    service, workspace = _service(tmp_path)
    config = {"paths": {"rules_manifest": str(tmp_path / "missing.json")}}

    exists, size = service.rules_manifest_state(config)

    assert exists is False
    assert size == 0


def test_log_run_writes_a_json_record_under_the_document_runs_dir(tmp_path):
    service, workspace = _service(tmp_path)
    config = {"paths": {}}
    path = service.log_run("doc1", config, tmp_path, "pipeline-prep", {"passed": True, "stages": []})
    assert path.parent == workspace.doc_root("doc1") / "runs"
    record = json.loads(path.read_text(encoding="utf-8"))
    assert record["command"] == "pipeline-prep"
    assert record["passed"] is True
    assert "timestamp" in record
    assert "git_commit" in record


def test_log_run_honors_configured_runs_dir_override(tmp_path):
    service, _ = _service(tmp_path)
    override_dir = tmp_path / "custom-runs"
    config = {"paths": {"runs_dir": str(override_dir)}}
    path = service.log_run("doc1", config, tmp_path, "pipeline-prep", {"passed": True})
    assert path.parent == override_dir


def test_list_runs_returns_empty_list_when_runs_dir_missing(tmp_path):
    service, _ = _service(tmp_path)
    assert service.list_runs("doc1", {"paths": {}}) == []


def test_list_runs_returns_records_most_recent_first(tmp_path):
    service, _ = _service(tmp_path)
    config = {"paths": {}}
    service.log_run("doc1", config, tmp_path, "pipeline-prep", {"n": 1})
    service.log_run("doc1", config, tmp_path, "pipeline-assemble", {"n": 2})
    records = service.list_runs("doc1", config)
    assert [r["n"] for r in records] == [2, 1]


def test_list_runs_respects_limit(tmp_path):
    service, _ = _service(tmp_path)
    config = {"paths": {}}
    for i in range(3):
        service.log_run("doc1", config, tmp_path, "pipeline-prep", {"n": i})
    assert len(service.list_runs("doc1", config, limit=2)) == 2


def test_list_runs_skips_malformed_json_files(tmp_path):
    service, workspace = _service(tmp_path)
    config = {"paths": {}}
    service.log_run("doc1", config, tmp_path, "pipeline-prep", {"n": 1})
    (workspace.doc_root("doc1") / "runs" / "broken.json").write_text("not json", encoding="utf-8")
    records = service.list_runs("doc1", config)
    assert len(records) == 1


def test_context_confirmed_lines_skips_sensitive_fields_and_includes_regular_ones(tmp_path):
    from docs.domain.models.template import Field, Template, Topic

    service, workspace = _service(tmp_path)
    template = Template(
        type="tesina",
        title="Tesina",
        context_schema={
            "topics": [
                Topic(
                    id="alumno",
                    title="Alumno",
                    fields=[
                        Field(key="nombre", label="Nombre", required=True),
                        Field(key="curp", label="CURP", required=False, sensitive=True),
                    ],
                )
            ]
        },
    )
    service.context_repository.write_topic("doc-1", template.context_schema.topics[0], {"nombre": "Ada", "curp": "AAAA000101HDFRRD01"})

    lines = service.context_confirmed_lines("doc-1", template)

    assert lines == ["Nombre: Ada"]


# --- Task 5: run_pipeline -----------------------------------------------
#
# NOTE on fixtures below: the plan's own draft `_template()`/`_pipeline_config()`
# were NOT copied verbatim. Two real discrepancies were found against the
# actual source and fixed here (see Task 5 final report for the full writeup):
#
# 1. `review_rules` (src/docs/domain/rules.py) unconditionally raises "error"
#    issues for a `Template` that lacks `paths.extracted_dir_policy`,
#    `preliminaries.roman_pagination`/`body_pagination_start`,
#    `format.page_margins_cm`, and an active `margins-2-5cm-non-cover`
#    advisor_override -- and for any declared `section_contracts` entry with
#    empty `required_content`. The plan's bare `_template()` (no extra
#    fields) and `_pipeline_config()` (no `type`/`title`, no matching
#    `section_contracts`) both fail `review_rules` unconditionally. Since
#    `review-rules` is a `fail_fast=True` prep stage, EVERY test that expects
#    stages after `review-rules` to run (build-sections, pack-context) would
#    never get there. Fixed by extending both fixtures with the same
#    review_rules-satisfying shape `tests/unit/domain/test_rules.py` already
#    uses for this purpose (`_valid_extra`), and by using
#    `Template.model_validate(...)` on a matching dict instead of the bare
#    constructor so `model_extra` is actually populated.
# 2. `DoctorService.run_doctor` (src/docs/application/doctor.py) appends
#    "pandoc" and "libreoffice" as `required=True` checks unconditionally
#    (not gated by `strict`, unlike "gh"). This host has pandoc but not
#    LibreOffice, so `doctor` would always fail regardless of config,
#    fail-fast-stopping every "prep" run at stage 1. Patched via the same
#    monkeypatch-the-resolver convention the plan itself already used for
#    "gh" (`shutil.which`), applied to `resolve_pandoc_executable`/
#    `resolve_libreoffice_executable` as imported into `docs.application.doctor`.
# 3. `del config["paths"]["context_dir"]` in the plan's fail-fast test does
#    NOT reproduce a KeyError: `run_doctor` reads `config["paths"].get(name)`
#    defensively, so a missing key is silently skipped (no check emitted at
#    all), and doctor would pass instead of failing. The actual, real way to
#    make the required `context_dir` check fail is to leave the configured
#    path pointing at a directory that is never created -- which is already
#    what happens if the `tmp_path / "context"` `.mkdir()` call is simply
#    omitted for that one test. The `del` line was removed accordingly.


def _valid_rules_extra() -> dict:
    """Matches tests/unit/domain/test_rules.py::_valid_extra() -- the minimal
    shape that makes review_rules() report zero "error" issues."""
    return {
        "preliminaries": {
            "roman_pagination": {"enabled": True},
            "body_pagination_start": {"section_id": "introduccion"},
        },
        "format": {
            "page_margins_cm": {
                "cover_policy": "preserve_template",
                "non_cover": {"top": 2.5, "right": 2.5, "bottom": 2.5, "left": 2.5},
            }
        },
        "advisor_overrides": [{"id": "margins-2-5cm-non-cover", "status": "active"}],
    }


def _template() -> Template:
    return Template.model_validate(
        {
            "type": "tesina",
            "title": "Tesina",
            "sections": [{"id": "introduccion", "title": "Introducción", "order": 1}],
            "section_contracts": {"introduccion": {"required_content": ["algo"]}},
            "paths": {"extracted_dir_policy": "rules_traceability_only"},
            **_valid_rules_extra(),
        }
    )


def _pipeline_config(tmp_path: Path) -> dict:
    return {
        "type": "tesina",
        "title": "Tesina",
        "paths": {
            "rules_manifest": str(tmp_path / "manual-rules.json"),
            "context_dir": str(tmp_path / "context"),
            "source_manifest": str(tmp_path / "source.json"),
            "issues_manifest": str(tmp_path / "issues.json"),
            "code_evidence_manifest": str(tmp_path / "code-evidence.json"),
            "fact_ledger": str(tmp_path / "00-fact-ledger.md"),
            "prompts_dir": str(tmp_path / "prompts"),
            "extracted_dir_policy": "rules_traceability_only",
        },
        "sections": [{"id": "introduccion", "title": "Introducción", "order": 1}],
        "section_contracts": {"introduccion": {"required_content": ["algo"]}},
        **_valid_rules_extra(),
        "evidence_sources": {},
        "privacy": {},
        "project": {},
    }


def _patch_doctor_tools(monkeypatch) -> None:
    """doctor's pandoc/libreoffice checks are required=True unconditionally
    (unlike gh, which is only required in --strict). Patched so `doctor`'s
    pass/fail in these tests reflects the fixture, not this host's toolchain.
    Task 2 (Slice 16, ToolResolverPort) moved DoctorService off the module-level
    resolve_pandoc_executable/resolve_libreoffice_executable imports it used to
    call directly, onto an injected ToolResolverPort (SystemToolResolverAdapter
    in these tests). The adapter still calls those same free functions, but now
    imports them into tool_resolver_adapter's namespace -- so the patch target
    moves there to keep intercepting the calls."""
    monkeypatch.setattr("docs.infrastructure.docx.tool_resolver_adapter.resolve_pandoc_executable", lambda paths: "pandoc")
    monkeypatch.setattr(
        "docs.infrastructure.docx.tool_resolver_adapter.resolve_libreoffice_executable", lambda paths: "soffice"
    )


def test_run_pipeline_prep_build_sections_succeeds_and_writes_the_section_file(tmp_path, monkeypatch):
    Path(tmp_path / "context").mkdir()
    service, _ = _service(tmp_path)
    _patch_doctor_tools(monkeypatch)
    monkeypatch.setattr("shutil.which", lambda name: None)  # gh unavailable -> collect-issues "omitido"
    summary = service.run_pipeline("doc1", _template(), _pipeline_config(tmp_path), "prep", repo_root=tmp_path)
    stage = next(s for s in summary["stages"] if s["stage"] == "build-sections")
    assert stage["ok"] is True
    assert stage["detail"] == "1 secciones"
    section_path = service.review_service.repository.section_path("doc1", 1, "introduccion")
    assert section_path.exists()
    assert "PENDIENTE: documentar algo con evidencia del ledger, contexto o fuentes." in section_path.read_text(
        encoding="utf-8"
    )


def test_run_pipeline_prep_runs_pack_context_after_build_sections(tmp_path, monkeypatch):
    # build-sections now succeeds (Task 5); this test only confirms the stage
    # ordering/continuation still holds, not a failure-recovery scenario.
    Path(tmp_path / "context").mkdir()
    service, _ = _service(tmp_path)
    _patch_doctor_tools(monkeypatch)
    monkeypatch.setattr("shutil.which", lambda name: None)
    summary = service.run_pipeline("doc1", _template(), _pipeline_config(tmp_path), "prep", repo_root=tmp_path)
    stage_names = [s["stage"] for s in summary["stages"]]
    assert "pack-context" in stage_names
    pack_context_stage = next(s for s in summary["stages"] if s["stage"] == "pack-context")
    assert pack_context_stage["ok"] is True


def test_run_pipeline_stops_at_first_fail_fast_failure(tmp_path):
    service, _ = _service(tmp_path)
    config = _pipeline_config(tmp_path)
    # `tmp_path / "context"` is intentionally never created: doctor's required
    # `context_dir` check (`path.exists() and path.is_dir()`) fails, which is
    # what actually reproduces the fail-fast-at-doctor scenario (see note above).
    summary = service.run_pipeline("doc1", _template(), config, "prep", repo_root=tmp_path)
    assert summary["passed"] is False
    assert summary["stages"][0]["stage"] == "doctor"
    assert len(summary["stages"]) == 1


def test_run_pipeline_writes_a_run_log_entry(tmp_path, monkeypatch):
    Path(tmp_path / "context").mkdir()
    service, workspace = _service(tmp_path)
    _patch_doctor_tools(monkeypatch)
    monkeypatch.setattr("shutil.which", lambda name: None)
    service.run_pipeline("doc1", _template(), _pipeline_config(tmp_path), "prep", repo_root=tmp_path)
    runs = service.list_runs("doc1", _pipeline_config(tmp_path))
    assert any(r["command"] == "pipeline-prep" for r in runs)


def test_run_pipeline_unknown_stage_set_raises_value_error(tmp_path):
    service, _ = _service(tmp_path)
    with pytest.raises(ValueError, match="Conjunto de etapas desconocido"):
        service.run_pipeline("doc1", _template(), _pipeline_config(tmp_path), "bogus", repo_root=tmp_path)


def test_run_pipeline_assemble_threads_custom_draft_name_to_audit_and_qa(tmp_path, monkeypatch):
    # Remediation (fresh-context review, WARNING): a custom
    # config["output"]["draft_name"] must reach format-audit-docx/qa-docx too
    # -- not just build-docx -- otherwise those stages look for the wrong
    # (missing, or worse, stale) hardcoded "tesina-draft.docx".
    class _FakeDocxRenderer:
        output_format = "docx"

        def stage_plan(self):
            return [("build-docx", True), ("format-audit-docx", True), ("qa-docx", True)]

        def build(self, doc_id, config, output=None):
            output_dir = Path(config["paths"]["output_draft_dir"])
            output_dir.mkdir(parents=True, exist_ok=True)
            name = config.get("output", {}).get("draft_name", "tesina-draft.docx")
            path = output or (output_dir / name)
            Document().save(path)
            return path

    service, _ = _service(tmp_path)
    monkeypatch.setattr(
        "docs.infrastructure.docx.libreoffice_qa_adapter.resolve_libreoffice_executable",
        lambda paths: None,
    )
    config = _pipeline_config(tmp_path)
    draft_dir = tmp_path / "draft"
    config["paths"]["output_draft_dir"] = str(draft_dir)
    config["paths"]["output_qa_dir"] = str(tmp_path / "qa")
    config["output"] = {"draft_name": "custom-draft.docx"}

    summary = service.run_pipeline(
        "doc1", _template(), config, "assemble", repo_root=tmp_path, renderer=_FakeDocxRenderer()
    )

    audit_stage = next(s for s in summary["stages"] if s["stage"] == "format-audit-docx")
    assert "No existe DOCX para auditar" not in audit_stage["detail"]
    # The correct custom-named file was produced; the stale default name never was.
    assert (draft_dir / "custom-draft.docx").exists()
    assert not (draft_dir / "tesina-draft.docx").exists()


def test_pipeline_and_renderer_resolve_draft_name_from_one_shared_default(tmp_path, monkeypatch):
    # D1 (tech-debt closeout): pipeline.py and docx_assembly.py each used to
    # declare their own "tesina-draft.docx" literal. Both must now resolve
    # the default from a single shared definition (docs.application.output_names)
    # -- patching that one place must change what BOTH modules resolve.
    monkeypatch.setattr("docs.application.output_names.DEFAULT_DRAFT_DOCX_NAME", "patched-draft.docx")
    service, workspace = _service(tmp_path)
    asset_service = AssetService(FilesystemAssetRepository(), workspace)
    renderer = DocxRendererAdapter(PythonDocxAssemblyAdapter(), asset_service, SystemToolResolverAdapter())

    assert service._resolve_draft_docx_name({}) == "patched-draft.docx"
    assert renderer._draft_docx_name({}) == "patched-draft.docx"


# --- Task 6: verify_all --------------------------------------------------
#
# NOTE: `verify_all` takes no `repo_root` parameter -- confirmed against the
# plan's Task 6 section ("Verbatim legacy reference: verify_all does not call
# collect_issues/collect_code_evidence/log_run, so it takes no repo_root
# parameter, unlike run_pipeline"), and reuses `_rules_manifest_state`/
# `resolve_normative_settings` from Task 5 rather than re-deriving them.


def test_verify_all_includes_review_rules_and_review_document_issues(tmp_path):
    Path(tmp_path / "context").mkdir()
    service, _ = _service(tmp_path)
    config = _pipeline_config(tmp_path)
    config["paths"]["output_draft_dir"] = str(tmp_path / "draft")  # no docx present -> docx-dependent checks skipped
    result = service.verify_all("doc1", _template(), config, strict=True)
    assert not result.passed  # missing rules_manifest -> review_rules error under strict


def test_verify_all_skips_docx_checks_when_no_draft_exists(tmp_path):
    Path(tmp_path / "context").mkdir()
    service, _ = _service(tmp_path)
    config = _pipeline_config(tmp_path)
    config["paths"]["output_draft_dir"] = str(tmp_path / "draft")
    result = service.verify_all("doc1", _template(), config, strict=False)
    assert not any(issue.code == "qa.failed" for issue in result.issues)


# Task 6's plan sketched a single loosely-asserting third test whose outcome
# depends on whether LibreOffice is installed in the execution environment.
# This is split into two variants, following test_libreoffice_qa_adapter.py's
# existing convention: one forces LibreOffice-unavailable deterministically
# via monkeypatch (works regardless of host toolchain, no skipif needed), one
# exercises the real success path and is skipif-skipped when LibreOffice is
# absent (as it is on this host).


def test_verify_all_appends_qa_failed_issue_when_libreoffice_unavailable(tmp_path, monkeypatch):
    Path(tmp_path / "context").mkdir()
    service, _ = _service(tmp_path)
    config = _pipeline_config(tmp_path)
    draft_dir = tmp_path / "draft"
    draft_dir.mkdir()
    config["paths"]["output_draft_dir"] = str(draft_dir)
    config["paths"]["output_qa_dir"] = str(tmp_path / "qa")
    docx_path = draft_dir / "tesina-draft.docx"
    Document().save(docx_path)
    monkeypatch.setattr(
        "docs.infrastructure.docx.libreoffice_qa_adapter.resolve_libreoffice_executable",
        lambda paths: None,
    )
    result = service.verify_all("doc1", _template(), config, strict=False)
    assert any(issue.code == "qa.failed" for issue in result.issues)


def test_verify_all_finds_docx_at_configured_draft_name(tmp_path, monkeypatch):
    # Remediation (fresh-context review, WARNING): verify_all's default-draft
    # lookup must honor config["output"]["draft_name"] too, not just the
    # hardcoded "tesina-draft.docx" -- otherwise a custom name makes verify_all
    # silently skip DOCX checks (candidate.exists() is False for the wrong name).
    Path(tmp_path / "context").mkdir()
    service, _ = _service(tmp_path)
    config = _pipeline_config(tmp_path)
    draft_dir = tmp_path / "draft"
    draft_dir.mkdir()
    config["paths"]["output_draft_dir"] = str(draft_dir)
    config["paths"]["output_qa_dir"] = str(tmp_path / "qa")
    config["output"] = {"draft_name": "custom-draft.docx"}
    Document().save(draft_dir / "custom-draft.docx")
    monkeypatch.setattr(
        "docs.infrastructure.docx.libreoffice_qa_adapter.resolve_libreoffice_executable",
        lambda paths: None,
    )

    result = service.verify_all("doc1", _template(), config, strict=False)

    # It found and audited the custom-named file (qa.failed comes from the
    # QA stage running against it, not from "no docx found at all").
    assert any(issue.code == "qa.failed" for issue in result.issues)


# --- Task 8.1: ingest stage_set wiring -----------------------------------


def _ingest_stage_config(workspace: Workspace, doc_id: str) -> dict:
    doc_root = workspace.doc_root(doc_id)
    return {
        "paths": {
            "inbox_dir": str(doc_root / "inbox"),
            "sections_dir": str(doc_root / "sections"),
            "context_dir": str(doc_root / "context"),
        },
    }


def test_run_pipeline_ingest_stage_set_writes_curated_index_without_touching_topic_qa_index(tmp_path):
    # Binding note carried from PR7's fresh review: `JsonContextRepository.
    # regenerate_index` (Topic/Q&A subsystem) and this module's new curated
    # progressive-disclosure index both target `context/`. Wiring must
    # namespace them under distinct filenames so neither writer clobbers
    # the other's most recent write.
    service, workspace = _service(tmp_path)
    doc_id = "doc1"
    topic = Topic(id="alumno", title="Alumno", required=True)
    schema = ContextSchema(topics=[topic])
    status = TopicStatus(id="alumno", title="Alumno", required=True, exists=False, missing=["(texto)"])
    service.context_repository.regenerate_index(doc_id, schema, [status])

    context_dir = workspace.doc_root(doc_id) / "context"
    topic_index_before = (context_dir / "index.md").read_text(encoding="utf-8")
    assert "Alumno" in topic_index_before  # sanity: the Topic/Q&A writer ran

    template = Template(type="doc", title="Doc")
    config = _ingest_stage_config(workspace, doc_id)

    summary = service.run_pipeline(doc_id, template, config, "ingest", repo_root=tmp_path)

    assert summary["passed"] is True
    topic_index_after = (context_dir / "index.md").read_text(encoding="utf-8")
    assert topic_index_after == topic_index_before  # untouched by the curation writer

    curated_index = (context_dir / "curated-index.md").read_text(encoding="utf-8")
    assert curated_index.startswith("# Context Index")
    assert curated_index != topic_index_after


@pytest.mark.skipif(not _HAS_LIBREOFFICE, reason="LibreOffice not installed")
def test_verify_all_completes_qa_without_qa_failed_when_libreoffice_available(tmp_path):
    Path(tmp_path / "context").mkdir()
    service, _ = _service(tmp_path)
    config = _pipeline_config(tmp_path)
    draft_dir = tmp_path / "draft"
    draft_dir.mkdir()
    config["paths"]["output_draft_dir"] = str(draft_dir)
    config["paths"]["output_qa_dir"] = str(tmp_path / "qa")
    docx_path = draft_dir / "tesina-draft.docx"
    Document().save(docx_path)
    result = service.verify_all("doc1", _template(), config, strict=False)
    assert not any(issue.code == "qa.failed" for issue in result.issues)
