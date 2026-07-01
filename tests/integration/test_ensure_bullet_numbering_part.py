from __future__ import annotations

import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

from docs.infrastructure.docx.python_docx_assembly_adapter import (
    ensure_bullet_numbering_part,
    set_bullet_numbering,
)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NUMBERING_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering"


def _docx_with_bulleted_paragraph(tmp_path: Path, num_id: int = 42) -> Path:
    document = Document()
    paragraph = document.add_paragraph("Item con vinieta")
    set_bullet_numbering(paragraph, num_id=num_id)
    path = tmp_path / "fixture.docx"
    document.save(path)
    return path


def _replace_zip_member(path: Path, member_name: str, content: str) -> None:
    # python-docx's default template already ships `word/numbering.xml` (for
    # its built-in list styles), so writing the poisoned payload requires
    # replacing that existing member rather than appending a duplicate entry
    # (zipfile raises "Duplicate name" under -W error for a naive append).
    with zipfile.ZipFile(path) as archive:
        members = {name: archive.read(name) for name in archive.namelist() if name != member_name}
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name, data in members.items():
            archive.writestr(name, data)
        archive.writestr(member_name, content)


def test_ensure_bullet_numbering_part_adds_numbering_xml_with_expected_num_id(tmp_path):
    path = _docx_with_bulleted_paragraph(tmp_path)
    ensure_bullet_numbering_part(path)

    with zipfile.ZipFile(path) as archive:
        assert "word/numbering.xml" in archive.namelist()
        numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
    root = ET.fromstring(numbering_xml)
    num = root.find(f".//{{{_W_NS}}}num[@{{{_W_NS}}}numId='42']")
    assert num is not None


def test_ensure_bullet_numbering_part_adds_relationship_and_content_type(tmp_path):
    path = _docx_with_bulleted_paragraph(tmp_path)
    ensure_bullet_numbering_part(path)

    with zipfile.ZipFile(path) as archive:
        rels_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        content_types_xml = archive.read("[Content_Types].xml").decode("utf-8")
    assert _NUMBERING_REL_TYPE in rels_xml
    assert "/word/numbering.xml" in content_types_xml


def test_ensure_bullet_numbering_part_is_a_noop_when_no_bulleted_paragraph_exists(tmp_path):
    # Note: python-docx's default template already ships a `word/numbering.xml`
    # (for its built-in list styles), so "no numbering.xml in the zip" is not
    # a valid no-op signal here. The real no-op invariant is that the file is
    # left byte-for-byte untouched: `ensure_bullet_numbering_part` extracts
    # into a temp dir and returns *before* ever rewriting `docx_path` when no
    # paragraph references `num_id`.
    document = Document()
    document.add_paragraph("Sin vinietas")
    path = tmp_path / "plain.docx"
    document.save(path)
    original_bytes = path.read_bytes()

    ensure_bullet_numbering_part(path)

    assert path.read_bytes() == original_bytes


def test_ensure_bullet_numbering_part_is_idempotent_across_repeated_calls(tmp_path):
    path = _docx_with_bulleted_paragraph(tmp_path)
    ensure_bullet_numbering_part(path)
    ensure_bullet_numbering_part(path)

    with zipfile.ZipFile(path) as archive:
        numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
        rels_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        content_types_xml = archive.read("[Content_Types].xml").decode("utf-8")
    root = ET.fromstring(numbering_xml)
    assert len(root.findall(f".//{{{_W_NS}}}num[@{{{_W_NS}}}numId='42']")) == 1
    assert rels_xml.count(_NUMBERING_REL_TYPE) == 1
    assert content_types_xml.count("/word/numbering.xml") == 1


def test_ensure_bullet_numbering_part_result_still_opens_via_python_docx(tmp_path):
    path = _docx_with_bulleted_paragraph(tmp_path)
    ensure_bullet_numbering_part(path)

    # Round-trip check: the rewritten zip must still be a valid .docx.
    reopened = Document(str(path))
    assert any(p.text == "Item con vinieta" for p in reopened.paragraphs)


def test_ensure_bullet_numbering_part_rejects_malicious_entity_expansion_in_existing_numbering_xml(tmp_path):
    # Hardening regression test for Design Decision 5.1 (defusedxml):
    # a pre-existing numbering.xml carrying a billion-laughs-style DTD
    # entity must be rejected by safe_parse, not silently expanded.
    path = _docx_with_bulleted_paragraph(tmp_path)
    _replace_zip_member(
        path,
        "word/numbering.xml",
        '<?xml version="1.0"?>'
        '<!DOCTYPE numbering [<!ENTITY lol "lol">]>'
        f'<w:numbering xmlns:w="{_W_NS}">&lol;</w:numbering>',
    )
    import pytest
    from defusedxml.common import DefusedXmlException

    with pytest.raises(DefusedXmlException):
        ensure_bullet_numbering_part(path)
