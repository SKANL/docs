# tests/integration/test_libreoffice_qa_adapter.py
from __future__ import annotations

import shutil
from pathlib import Path

import pytest
from docx import Document

from docs.infrastructure.docx.libreoffice_qa_adapter import (
    LibreOfficeQaAdapter,
    resolve_libreoffice_executable,
)

_HAS_LIBREOFFICE = shutil.which("soffice") is not None or shutil.which("libreoffice") is not None


def test_resolve_libreoffice_executable_prefers_path_lookup():
    if not _HAS_LIBREOFFICE:
        pytest.skip("LibreOffice not installed")
    assert resolve_libreoffice_executable({}) is not None


def test_resolve_libreoffice_executable_falls_back_to_configured_bin(tmp_path, monkeypatch):
    monkeypatch.setattr(shutil, "which", lambda name: None)
    fake_bin = tmp_path / "soffice.exe"
    fake_bin.write_text("", encoding="utf-8")
    assert resolve_libreoffice_executable({"libreoffice_bin": str(fake_bin)}) == str(fake_bin)


def test_resolve_libreoffice_executable_returns_none_when_nothing_matches(monkeypatch):
    monkeypatch.setattr(shutil, "which", lambda name: None)
    assert resolve_libreoffice_executable({}) is None


@pytest.mark.skipif(not _HAS_LIBREOFFICE, reason="LibreOffice not installed")
def test_render_docx_to_pdf_produces_a_non_empty_pdf(tmp_path):
    docx_path = tmp_path / "doc.docx"
    Document().save(docx_path)
    output_dir = tmp_path / "out"
    output_dir.mkdir()

    pdf_path = LibreOfficeQaAdapter().render_docx_to_pdf({"paths": {}}, docx_path, output_dir)

    assert pdf_path == output_dir / "doc.pdf"
    assert pdf_path.stat().st_size > 0


def test_render_docx_to_pdf_raises_when_libreoffice_unavailable(tmp_path, monkeypatch):
    monkeypatch.setattr(
        "docs.infrastructure.docx.libreoffice_qa_adapter.resolve_libreoffice_executable", lambda paths: None
    )
    docx_path = tmp_path / "doc.docx"
    Document().save(docx_path)
    output_dir = tmp_path / "out"
    output_dir.mkdir()

    with pytest.raises(RuntimeError, match="LibreOffice/soffice no está disponible"):
        LibreOfficeQaAdapter().render_docx_to_pdf({"paths": {}}, docx_path, output_dir)


def test_run_documents_audits_returns_empty_list_when_disabled(tmp_path):
    result = LibreOfficeQaAdapter().run_documents_audits(
        {"documents_tools": {"enabled": False}}, tmp_path / "doc.docx", tmp_path
    )
    assert result == []


def test_run_documents_audits_marks_missing_scripts_dir_as_not_found_non_strict(tmp_path):
    result = LibreOfficeQaAdapter().run_documents_audits({}, tmp_path / "doc.docx", tmp_path, strict=False)
    assert len(result) == 4
    assert all(item["ok"] is True and item["stderr"] == "script no encontrado" for item in result)


def test_run_documents_audits_marks_missing_scripts_dir_as_failing_strict(tmp_path):
    result = LibreOfficeQaAdapter().run_documents_audits({}, tmp_path / "doc.docx", tmp_path, strict=True)
    assert all(item["ok"] is False for item in result)


def test_run_documents_audits_runs_a_real_configured_script_and_captures_output(tmp_path):
    scripts_dir = tmp_path / "scripts"
    scripts_dir.mkdir()
    script = scripts_dir / "heading_audit.py"
    script.write_text("print('ok from heading_audit')\n", encoding="utf-8")
    output_dir = tmp_path / "out"
    output_dir.mkdir()
    docx_path = tmp_path / "doc.docx"
    Document().save(docx_path)

    result = LibreOfficeQaAdapter().run_documents_audits(
        {"paths": {"documents_scripts_dir": str(scripts_dir)}}, docx_path, output_dir
    )

    heading_result = next(item for item in result if item["name"] == "heading_audit.py")
    assert heading_result["ok"] is True
    assert "ok from heading_audit" in heading_result["stdout"]
    assert Path(heading_result["report"]).exists()
