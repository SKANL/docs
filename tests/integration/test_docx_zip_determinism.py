# tests/integration/test_docx_zip_determinism.py
from __future__ import annotations

import json
import shutil
import struct
import subprocess
import time as time_module
import zipfile
import zlib
from pathlib import Path

import pytest
from docx import Document

from docs.application.asset import AssetService
from docs.application.docx_assembly import DocxRendererAdapter
from docs.domain.workspace import Workspace
from docs.infrastructure.docx.deterministic_zip import SENTINEL_CORE_XML_DATETIME, SENTINEL_DATE_TIME
from docs.infrastructure.docx.python_docx_assembly_adapter import PythonDocxAssemblyAdapter
from docs.infrastructure.docx.tool_resolver_adapter import SystemToolResolverAdapter
from docs.infrastructure.persistence.filesystem_asset_repository import FilesystemAssetRepository


def _save_body_docx(tmp_path: Path, name: str = "body.docx") -> Path:
    document = Document()
    document.add_heading("Introduccion", level=1)
    document.add_paragraph("Texto de cuerpo.")
    path = tmp_path / name
    document.save(path)
    return path


def test_assemble_output_is_byte_identical_across_a_2_second_dos_timestamp_boundary(tmp_path, monkeypatch):
    # Root cause: python-docx's zip writer (`docx.opc.phys_pkg`) calls
    # `ZipFile.writestr(arcname, blob)` with no explicit `ZipInfo`, so
    # zipfile stamps every entry's `date_time` with `time.localtime(time.time())`
    # at the moment of the call -- 2-second DOS timestamp granularity. Two
    # builds of an otherwise identical document that straddle a 2-second
    # boundary produce a byte-level diff purely from zip metadata, violating
    # this harness's same-inputs -> byte-identical-outputs invariant. The
    # fixture body here (heading + plain paragraph, no bulleted list) mirrors
    # the real failure: no `w:numId` reference means
    # `ensure_bullet_numbering_part` takes its early-return, no-op path, so
    # nothing after `main.save(output_docx)` would otherwise touch the file.
    body = _save_body_docx(tmp_path)
    real_time = time_module.time

    def build(fixed_now: float, name: str) -> bytes:
        monkeypatch.setattr("time.time", lambda: fixed_now)
        try:
            output = tmp_path / name
            PythonDocxAssemblyAdapter().assemble(
                {}, body, output, cover_asset_path=None, embed_front_paths=[], embed_back_paths=[]
            )
            return output.read_bytes()
        finally:
            monkeypatch.setattr("time.time", real_time)

    t0 = real_time()
    first = build(t0, "first.docx")
    second = build(t0 + 2.5, "second.docx")  # straddles a 2-second DOS timestamp boundary

    assert first == second


def test_assemble_zip_entries_use_a_fixed_sentinel_timestamp(tmp_path, monkeypatch):
    body = _save_body_docx(tmp_path)
    output = tmp_path / "out.docx"
    PythonDocxAssemblyAdapter().assemble(
        {}, body, output, cover_asset_path=None, embed_front_paths=[], embed_back_paths=[]
    )
    with zipfile.ZipFile(output) as archive:
        timestamps = {info.date_time for info in archive.infolist()}
    assert timestamps == {(1980, 1, 1, 0, 0, 0)}


def test_assemble_embed_branch_normalizes_zip_timestamps_and_is_deterministic(tmp_path):
    # Coverage gap closed here (fresh-context review SUGGESTION, round 2): the
    # only two embed-branch tests in test_python_docx_assembly_adapter.py force
    # docxcompose to be unimportable, so they only exercise the RuntimeError
    # path. Nothing locked in the embed-branch SUCCESS path -- real
    # docxcompose.Composer, non-empty embed_front_paths/embed_back_paths,
    # composer.save(), then normalize_docx_zip_timestamps -- ever running.
    # docxcompose is a declared, always-installed dependency (pyproject.toml),
    # unlike pandoc/java, so no availability skip is needed here.
    body = _save_body_docx(tmp_path)
    front = Document()
    front.add_paragraph("FRONT MATTER MARKER")
    front_path = tmp_path / "front.docx"
    front.save(front_path)

    back = Document()
    back.add_paragraph("BACK MATTER MARKER")
    back_path = tmp_path / "back.docx"
    back.save(back_path)

    def build(name: str) -> bytes:
        output = tmp_path / name
        PythonDocxAssemblyAdapter().assemble(
            {},
            body,
            output,
            cover_asset_path=None,
            embed_front_paths=[front_path],
            embed_back_paths=[back_path],
        )
        return output.read_bytes()

    first = build("first.docx")
    second = build("second.docx")
    assert first == second

    with zipfile.ZipFile(tmp_path / "first.docx") as archive:
        timestamps = {info.date_time for info in archive.infolist()}
    assert timestamps == {SENTINEL_DATE_TIME}


# --- render_pandoc (body docx written directly by the pandoc subprocess) ------


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_render_pandoc_output_is_deterministic_by_construction(tmp_path):
    # Root cause: pandoc is an external subprocess that writes the .docx
    # itself -- unlike the three python-docx write sites this module already
    # normalizes, nothing ever called `normalize_docx_zip_timestamps` on
    # `render_pandoc`'s output. pandoc stamps every zip entry's `date_time`
    # with the wall clock AND writes a real `dcterms:created`/
    # `dcterms:modified` value into docProps/core.xml, so the resulting body
    # .docx violated this harness's same-inputs -> byte-identical-outputs
    # invariant on both zip metadata and payload bytes.
    markdown = tmp_path / "section.md"
    markdown.write_text("# Titulo\n\nCuerpo del texto de prueba.\n", encoding="utf-8")
    output = tmp_path / "body.docx"
    PythonDocxAssemblyAdapter().render_pandoc(shutil.which("pandoc"), [markdown], output)

    with zipfile.ZipFile(output) as archive:
        timestamps = {info.date_time for info in archive.infolist()}
        core_xml = archive.read("docProps/core.xml").decode("utf-8")

    assert timestamps == {SENTINEL_DATE_TIME}
    assert core_xml.count(SENTINEL_CORE_XML_DATETIME) == 2  # dcterms:created + dcterms:modified


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_render_pandoc_output_is_byte_identical_across_a_real_wall_clock_gap(tmp_path):
    # pandoc's wall clock is read inside a separate subprocess, so unlike the
    # python-docx write sites (which monkeypatch `time.time`), it cannot be
    # faked from this process. This is the one test in the suite that uses a
    # real sleep to prove byte-identity across an actual elapsed time gap.
    markdown = tmp_path / "section.md"
    markdown.write_text("# Titulo\n\nCuerpo del texto de prueba.\n", encoding="utf-8")
    pandoc_path = shutil.which("pandoc")

    first = tmp_path / "first.docx"
    PythonDocxAssemblyAdapter().render_pandoc(pandoc_path, [markdown], first)
    time_module.sleep(2.1)
    second = tmp_path / "second.docx"
    PythonDocxAssemblyAdapter().render_pandoc(pandoc_path, [markdown], second)

    assert first.read_bytes() == second.read_bytes()


# --- bound-figure embedding: byte-identical rebuild (S4, ADR-5/ADR-7) ---------


def _png_chunk(tag: bytes, data: bytes) -> bytes:
    return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", zlib.crc32(tag + data))


def _solid_png(width: int, height: int) -> bytes:
    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)  # 8-bit RGB
    raw = b"".join(b"\x00" + bytes([180, 180, 180] * width) for _ in range(height))
    idat = zlib.compress(raw)
    return b"\x89PNG\r\n\x1a\n" + _png_chunk(b"IHDR", ihdr) + _png_chunk(b"IDAT", idat) + _png_chunk(b"IEND", b"")


@pytest.mark.skipif(shutil.which("pandoc") is None, reason="pandoc not installed")
def test_build_with_bound_figure_is_byte_identical_across_repeated_runs(tmp_path):
    # Characterization, not new behavior (tasks.md S4 4.8): embedding a bound
    # figure introduces no clock/random data, and the absolute path in the
    # intermediate markdown is only pandoc's read handle -- it never enters
    # the output bytes (design.md ADR-5). This must already pass once 4.7's
    # wiring is correct; a failure here is a real determinism bug.
    workspace = Workspace(documents_dir=tmp_path / "documents", templates_dir=tmp_path / "templates")
    asset_service = AssetService(FilesystemAssetRepository(), workspace)
    assets_dir = workspace.assets_dir("doc-1")
    figures_dir = assets_dir / "figures"
    figures_dir.mkdir(parents=True)
    (figures_dir / "fig-abc12345.png").write_bytes(_solid_png(150, 100))

    sections_dir = tmp_path / "sections"
    sections_dir.mkdir()
    (sections_dir / "001-resumen.md").write_text(
        "# Resumen\n\n[[figure:organigrama]] Organigrama del equipo.\n", encoding="utf-8"
    )
    (sections_dir / "figure-catalog.json").write_text(
        json.dumps(
            {
                "figures": [
                    {
                        "id": "fig-abc12345",
                        "sha256": "0" * 64,
                        "width_px": 150,
                        "height_px": 100,
                        "origin_relative_path": "assets/figures/fig-abc12345.png",
                        "caption": "Organigrama del equipo",
                        "source_role": "evidence",
                        "origin_kind": "standalone",
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    (sections_dir / "figure-bindings.json").write_text(
        json.dumps({"schema": 1, "bindings": {"organigrama": "fig-abc12345"}}), encoding="utf-8"
    )

    seed = tmp_path / "_seed_template.md"
    seed.write_text("Plantilla.\n", encoding="utf-8")
    template = tmp_path / "template.docx"
    subprocess.run([shutil.which("pandoc"), str(seed), "-o", str(template)], check=True)

    config = {
        "sections": [{"id": "resumen", "order": 1}],
        "paths": {
            "sections_dir": str(sections_dir),
            "output_draft_dir": str(tmp_path / "draft"),
            "template_docx": str(template),
            "assets_dir": str(assets_dir),
        },
    }
    service = DocxRendererAdapter(PythonDocxAssemblyAdapter(), asset_service, SystemToolResolverAdapter())

    first = service.build("doc-1", config, output=tmp_path / "draft" / "first.docx")
    second = service.build("doc-1", config, output=tmp_path / "draft" / "second.docx")

    assert first.read_bytes() == second.read_bytes()
    with zipfile.ZipFile(first) as archive:
        assert any(name.startswith("word/media/") for name in archive.namelist())
