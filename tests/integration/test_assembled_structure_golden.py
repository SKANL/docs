# tests/integration/test_assembled_structure_golden.py
"""One whole document, built end to end, pinned by its STRUCTURE.

Everything else in this suite checks assembly one property at a time: this
figure is centered, that table repeats its header, these bytes are identical
across two runs. Nothing looked at a finished document as a whole -- the two
acceptance tests stop at `prep` and never build a `.docx` at all. So a
regression that leaves every individual property intact while getting their
COMBINATION wrong (a section break lost, numbering restarting in the wrong
place, a cross-reference marker surviving into the output) had nothing to
fail against.

Why a structural golden and not a byte golden: a committed `.docx` breaks on
every python-docx or pandoc bump, which teaches people to regenerate goldens
without reading them -- and a golden nobody reads is worse than no golden.
The structure summary below is stable across those bumps and still changes
the moment the assembler's real behaviour does.

Regenerate deliberately, never reflexively:

    UPDATE_STRUCTURE_GOLDEN=1 uv run pytest tests/integration/test_assembled_structure_golden.py

then READ the diff. It is the assembler telling you what you changed.
"""
from __future__ import annotations

import os
import shutil
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from docs.application.asset import AssetService
from docs.application.docx_assembly import DocxRendererAdapter
from docs.domain.workspace import Workspace
from docs.infrastructure.docx.python_docx_assembly_adapter import PythonDocxAssemblyAdapter
from docs.infrastructure.docx.tool_resolver_adapter import SystemToolResolverAdapter
from docs.infrastructure.persistence.filesystem_asset_repository import FilesystemAssetRepository

GOLDEN = Path(__file__).resolve().parent / "goldens" / "assembled-structure.txt"
_HAS_PANDOC = shutil.which("pandoc") is not None

_SECTIONS = {
    "001-introduccion.md": (
        "# Introducción\n\n"
        "Este documento describe el arnés. La arquitectura aparece en "
        "[[figure:arquitectura]], y los resultados en [[table:resumen]].\n\n"
        "## Alcance\n\n"
        "Cubre la generación determinista de documentos.\n"
    ),
    "002-resultados.md": (
        "# Resultados\n\n"
        "El detalle completo está en [[table:resumen]].\n\n"
        "| Métrica | Valor |\n"
        "|---|---|\n"
        "| Cobertura | 96% |\n"
        "| Tests | 1566 |\n\n"
        "Tabla 1. Resumen de métricas.\n\n"
        "- primer punto\n"
        "- segundo punto\n"
    ),
}


def _structure(path: Path) -> str:
    """A stable, readable summary of what the assembler actually produced."""
    document = Document(str(path))
    lines: list[str] = []

    lines.append("# sections")
    for index, section in enumerate(document.sections):
        pg = section._sectPr.find(qn("w:pgNumType"))
        start = pg.get(qn("w:start")) if pg is not None else None
        fmt = pg.get(qn("w:fmt")) if pg is not None else None
        lines.append(f"  [{index}] page-numbering start={start} format={fmt}")

    lines.append("# outline")
    for paragraph in document.paragraphs:
        style = paragraph.style.name if paragraph.style is not None else "?"
        text = paragraph.text.strip()
        if style.startswith("Heading"):
            lines.append(f"  {style}: {text}")

    lines.append("# paragraph styles in use")
    counts: dict[str, int] = {}
    for paragraph in document.paragraphs:
        style = paragraph.style.name if paragraph.style is not None else "?"
        counts[style] = counts.get(style, 0) + 1
    for style in sorted(counts):
        lines.append(f"  {style}: {counts[style]}")

    lines.append("# tables")
    for index, table in enumerate(document.tables):
        lines.append(f"  [{index}] rows={len(table.rows)} cols={len(table.columns)}")

    lines.append("# resolved markers")
    body = "\n".join(p.text for p in document.paragraphs)
    lines.append(f"  unresolved '[[' markers: {body.count('[[')}")
    lines.append(f"  figure captions: {sorted({w for w in body.split() if w == 'Figura'})}")

    lines.append("# fields")
    xml = document.element.body.xml
    lines.append(f"  TOC field present: {'TOC' in xml}")
    lines.append(f"  PAGE field present: {'PAGE' in xml}")

    return "\n".join(lines) + "\n"


def _build(root: Path) -> Path:
    """Assemble the fixture document under `root` and return the artifact."""
    workspace = Workspace(documents_dir=root / "documents", templates_dir=root / "templates")
    sections_dir = workspace.doc_root("doc-1") / "sections"
    sections_dir.mkdir(parents=True)
    for name, text in _SECTIONS.items():
        (sections_dir / name).write_text(text, encoding="utf-8")

    output_dir = workspace.doc_root("doc-1") / "output" / "draft"
    output_dir.mkdir(parents=True)
    config = {
        "doc_id": "doc-1",
        "sections": [
            {"id": "introduccion", "order": 1},
            {"id": "resultados", "order": 2},
        ],
        "paths": {
            "sections_dir": str(sections_dir),
            "output_draft_dir": str(output_dir),
            "assets_dir": str(workspace.assets_dir("doc-1")),
        },
        "format": {"page_size": "letter"},
        "structure": [{"type": "toc"}, {"type": "sections"}],
    }
    renderer = DocxRendererAdapter(
        PythonDocxAssemblyAdapter(),
        AssetService(FilesystemAssetRepository(), workspace),
        SystemToolResolverAdapter(),
    )
    built = renderer.build("doc-1", config)
    assert built is not None
    return built


@pytest.fixture
def built_document(tmp_path: Path) -> Path:
    return _build(tmp_path)


@pytest.mark.skipif(not _HAS_PANDOC, reason="pandoc not installed")
def test_assembled_document_structure_matches_the_golden(built_document: Path):
    produced = _structure(built_document)

    if os.environ.get("UPDATE_STRUCTURE_GOLDEN"):
        GOLDEN.parent.mkdir(parents=True, exist_ok=True)
        GOLDEN.write_text(produced, encoding="utf-8")
        pytest.skip("golden regenerado — leé el diff antes de commitear")

    assert GOLDEN.is_file(), (
        f"falta el golden en {GOLDEN}. Generalo con "
        f"UPDATE_STRUCTURE_GOLDEN=1 y REVISÁ el contenido antes de commitear."
    )
    expected = GOLDEN.read_text(encoding="utf-8")
    assert produced == expected, (
        "la estructura del documento ensamblado cambió.\n\n"
        "Si el cambio es deliberado, regeneralo con UPDATE_STRUCTURE_GOLDEN=1 "
        "y leé el diff: es el ensamblador contándote qué cambiaste.\n\n"
        f"--- golden\n{expected}\n--- producido\n{produced}"
    )


@pytest.mark.skipif(not _HAS_PANDOC, reason="pandoc not installed")
def test_the_golden_covers_what_it_claims_to_cover(built_document: Path):
    # A golden that summarises nothing passes forever. This pins the summary
    # itself: if `_structure` stops reporting a dimension, the golden silently
    # stops guarding it.
    produced = _structure(built_document)
    for heading in ("# sections", "# outline", "# paragraph styles in use",
                    "# tables", "# resolved markers", "# fields"):
        assert heading in produced, heading
    assert "Heading 1: Introducción" in produced, produced
    assert "unresolved '[[' markers: 0" in produced, produced


@pytest.mark.skipif(not _HAS_PANDOC, reason="pandoc not installed")
def test_this_richer_fixture_is_still_byte_identical_across_builds(tmp_path: Path):
    # The byte-identity tests elsewhere use a heading plus a plain paragraph.
    # This fixture adds a TOC field, a table, a bulleted list and resolved
    # cross-references -- every construct that writes its own XML and could
    # therefore smuggle in a wall-clock value.
    first = _build(tmp_path / "one")
    second = _build(tmp_path / "two")

    assert first.read_bytes() == second.read_bytes()
