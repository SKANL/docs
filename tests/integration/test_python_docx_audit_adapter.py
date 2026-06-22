# tests/integration/test_python_docx_audit_adapter.py
from __future__ import annotations

from pathlib import Path
from typing import Callable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from docs.infrastructure.docx.python_docx_audit_adapter import PythonDocxAuditAdapter


def _save_minimal_docx(tmp_path: Path, build: Callable[[Document], None]) -> Path:
    document = Document()
    build(document)
    path = tmp_path / "fixture.docx"
    document.save(path)
    return path


# --- heading casing / numbering ---------------------------------------------


def test_audit_warns_on_lowercase_heading_1(tmp_path):
    def build(document):
        document.add_heading("Capitulo Uno", level=1)

    path = _save_minimal_docx(tmp_path, build)
    issues = PythonDocxAuditAdapter().audit(path, {}, strict=False)
    assert any("mayúsculas sostenidas" in issue.message for issue in issues)
    assert all(issue.severity == "warning" for issue in issues if "mayúsculas sostenidas" in issue.message)


def test_audit_no_warning_for_uppercase_heading_1(tmp_path):
    def build(document):
        document.add_heading("CAPITULO UNO", level=1)

    path = _save_minimal_docx(tmp_path, build)
    issues = PythonDocxAuditAdapter().audit(path, {}, strict=False)
    assert not any("mayúsculas sostenidas" in issue.message for issue in issues)


def test_audit_warns_on_manually_numbered_heading_1(tmp_path):
    def build(document):
        document.add_heading("1.1 INTRODUCCION", level=1)

    path = _save_minimal_docx(tmp_path, build)
    issues = PythonDocxAuditAdapter().audit(path, {}, strict=False)
    assert any("numerado manualmente" in issue.message for issue in issues)


def test_audit_no_numbering_warning_for_plain_heading_1(tmp_path):
    def build(document):
        document.add_heading("INTRODUCCION", level=1)

    path = _save_minimal_docx(tmp_path, build)
    issues = PythonDocxAuditAdapter().audit(path, {}, strict=False)
    assert not any("numerado manualmente" in issue.message for issue in issues)


# --- table borders / shading -------------------------------------------------


def test_audit_flags_table_with_vertical_borders(tmp_path):
    def build(document):
        table = document.add_table(rows=1, cols=2)
        tbl_pr = table._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        borders.append(left)
        tbl_pr.append(borders)

    path = _save_minimal_docx(tmp_path, build)
    issues = PythonDocxAuditAdapter().audit(path, {}, strict=False)
    assert any("bordes verticales" in issue.message for issue in issues)
    assert any(issue.severity == "error" for issue in issues if "bordes verticales" in issue.message)


def test_audit_no_table_warning_for_table_without_vertical_borders_or_shading(tmp_path):
    def build(document):
        document.add_table(rows=1, cols=2)

    path = _save_minimal_docx(tmp_path, build)
    issues = PythonDocxAuditAdapter().audit(path, {}, strict=False)
    assert not any("bordes verticales" in issue.message for issue in issues)


# --- strict-mode section margin mismatch ------------------------------------


def test_audit_strict_flags_section_margin_mismatch(tmp_path):
    def build(document):
        document.add_heading("Introduccion", level=1)
        document.add_section()
        document.sections[1].top_margin = Cm(1.0)
        document.sections[1].right_margin = Cm(1.0)
        document.sections[1].bottom_margin = Cm(1.0)
        document.sections[1].left_margin = Cm(1.0)

    path = _save_minimal_docx(tmp_path, build)
    config = {"format": {"page_margins_cm": {"non_cover": {"top": 2.5, "right": 2.5, "bottom": 2.5, "left": 2.5}}}}
    issues = PythonDocxAuditAdapter().audit(path, config, strict=True)
    assert any("márgenes de 2.5 cm" in issue.message for issue in issues)
    assert any(issue.severity == "error" for issue in issues if "márgenes de 2.5 cm" in issue.message)


def test_audit_strict_does_not_flag_matching_section_margins(tmp_path):
    def build(document):
        document.add_heading("Introduccion", level=1)
        document.add_section()
        document.sections[1].top_margin = Cm(2.5)
        document.sections[1].right_margin = Cm(2.5)
        document.sections[1].bottom_margin = Cm(2.5)
        document.sections[1].left_margin = Cm(2.5)

    path = _save_minimal_docx(tmp_path, build)
    config = {"format": {"page_margins_cm": {"non_cover": {"top": 2.5, "right": 2.5, "bottom": 2.5, "left": 2.5}}}}
    issues = PythonDocxAuditAdapter().audit(path, config, strict=True)
    assert not any("márgenes de 2.5 cm" in issue.message for issue in issues)


def test_audit_non_strict_does_not_check_margins(tmp_path):
    def build(document):
        document.add_heading("Introduccion", level=1)
        document.add_section()
        document.sections[1].top_margin = Cm(1.0)
        document.sections[1].right_margin = Cm(1.0)
        document.sections[1].bottom_margin = Cm(1.0)
        document.sections[1].left_margin = Cm(1.0)

    path = _save_minimal_docx(tmp_path, build)
    config = {"format": {"page_margins_cm": {"non_cover": {"top": 2.5, "right": 2.5, "bottom": 2.5, "left": 2.5}}}}
    issues = PythonDocxAuditAdapter().audit(path, config, strict=False)
    assert not any("márgenes de 2.5 cm" in issue.message for issue in issues)


# --- figure caption -----------------------------------------------------------


def test_audit_flags_image_without_figura_caption(tmp_path):
    def build(document):
        document.add_heading("Introduccion", level=1)
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        drawing = OxmlElement("w:drawing")
        run._r.append(drawing)
        document.add_paragraph("Texto sin caption.")

    path = _save_minimal_docx(tmp_path, build)
    issues = PythonDocxAuditAdapter().audit(path, {}, strict=False)
    assert any("Figura" in issue.message for issue in issues)
    assert any(issue.severity == "warning" for issue in issues if "Figura" in issue.message)


def test_audit_no_figure_warning_when_caption_follows(tmp_path):
    def build(document):
        document.add_heading("Introduccion", level=1)
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        drawing = OxmlElement("w:drawing")
        run._r.append(drawing)
        document.add_paragraph("Figura 1. Descripcion de la figura.")

    path = _save_minimal_docx(tmp_path, build)
    issues = PythonDocxAuditAdapter().audit(path, {}, strict=False)
    assert not any("Figura detectada sin caption" in issue.message for issue in issues)


# --- list_parts / read_xml ---------------------------------------------------


def test_list_parts_returns_xml_parts_with_prefix(tmp_path):
    path = _save_minimal_docx(tmp_path, lambda d: d.add_paragraph("x"))
    parts = PythonDocxAuditAdapter().list_parts(path, "word/")
    assert "word/document.xml" in parts


def test_list_parts_excludes_non_matching_prefix(tmp_path):
    path = _save_minimal_docx(tmp_path, lambda d: d.add_paragraph("x"))
    parts = PythonDocxAuditAdapter().list_parts(path, "word/footer")
    assert "word/document.xml" not in parts


def test_read_xml_returns_decoded_content_for_known_part(tmp_path):
    path = _save_minimal_docx(tmp_path, lambda d: d.add_paragraph("contenido unico"))
    xml = PythonDocxAuditAdapter().read_xml(path, "word/document.xml")
    assert "contenido" in xml


def test_read_xml_returns_empty_string_for_unknown_part(tmp_path):
    path = _save_minimal_docx(tmp_path, lambda d: d.add_paragraph("x"))
    assert PythonDocxAuditAdapter().read_xml(path, "word/missing.xml") == ""
