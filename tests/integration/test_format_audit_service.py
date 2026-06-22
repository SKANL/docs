# tests/integration/test_format_audit_service.py
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docs.application.format_audit import FormatAuditService
from docs.infrastructure.docx.python_docx_audit_adapter import PythonDocxAuditAdapter


@pytest.fixture
def service() -> FormatAuditService:
    return FormatAuditService(PythonDocxAuditAdapter())


def test_audit_format_raises_filenotfound_when_docx_missing(tmp_path: Path, service: FormatAuditService) -> None:
    with pytest.raises(FileNotFoundError):
        service.audit_format(tmp_path / "missing.docx", {})


def test_audit_format_returns_clean_result_for_compliant_docx(tmp_path: Path, service: FormatAuditService) -> None:
    document = Document()
    document.add_heading("INTRODUCCION", level=1)
    document.add_paragraph("Texto de cuerpo sin hallazgos esperables.")
    path = tmp_path / "ok.docx"
    document.save(path)

    result = service.audit_format(path, {}, strict=False)

    assert result.passed is True


def test_audit_format_returns_result_with_issues_for_lowercase_heading(
    tmp_path: Path, service: FormatAuditService
) -> None:
    document = Document()
    document.add_heading("Introduccion", level=1)
    path = tmp_path / "bad.docx"
    document.save(path)

    result = service.audit_format(path, {}, strict=False)

    assert result.issues
    assert any("mayúsculas sostenidas" in issue.message for issue in result.issues)
