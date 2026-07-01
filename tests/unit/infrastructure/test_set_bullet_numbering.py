from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn

from docs.infrastructure.docx.python_docx_assembly_adapter import set_bullet_numbering


def test_set_bullet_numbering_adds_num_pr_with_default_num_id():
    document = Document()
    paragraph = document.add_paragraph("Item")
    set_bullet_numbering(paragraph)
    p_pr = paragraph._p.pPr
    num_pr = p_pr.find(qn("w:numPr"))
    assert num_pr is not None
    assert num_pr.find(qn("w:ilvl")).get(qn("w:val")) == "0"
    assert num_pr.find(qn("w:numId")).get(qn("w:val")) == "42"


def test_set_bullet_numbering_honors_custom_num_id():
    document = Document()
    paragraph = document.add_paragraph("Item")
    set_bullet_numbering(paragraph, num_id=7)
    num_pr = paragraph._p.pPr.find(qn("w:numPr"))
    assert num_pr.find(qn("w:numId")).get(qn("w:val")) == "7"


def test_set_bullet_numbering_is_idempotent_not_duplicating_elements():
    document = Document()
    paragraph = document.add_paragraph("Item")
    set_bullet_numbering(paragraph)
    set_bullet_numbering(paragraph)
    p_pr = paragraph._p.pPr
    assert len(p_pr.findall(qn("w:numPr"))) == 1
    num_pr = p_pr.find(qn("w:numPr"))
    assert len(num_pr.findall(qn("w:numId"))) == 1


def test_set_bullet_numbering_round_trips_through_save_and_reopen(tmp_path):
    document = Document()
    paragraph = document.add_paragraph("Item")
    set_bullet_numbering(paragraph)
    path = tmp_path / "fixture.docx"
    document.save(path)

    reopened = Document(str(path))
    target = next(p for p in reopened.paragraphs if p.text == "Item")
    num_pr = target._p.pPr.find(qn("w:numPr"))
    assert num_pr.find(qn("w:numId")).get(qn("w:val")) == "42"
