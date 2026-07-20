import base64
import io

from docx import Document
from docx.oxml.ns import qn

from docs.infrastructure.docx.python_docx_assembly_adapter import PythonDocxAssemblyAdapter

# 1x1 pixel PNG, same fixture used across the ingest tests.
_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY"
    "42YAAAAASUVORK5CYII="
)


def test_transfer_body_tables_copies_cell_text_into_new_table():
    adapter = PythonDocxAssemblyAdapter()
    cover = Document()
    body = Document()
    table = body.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "a"
    table.cell(0, 1).text = "b"
    table.cell(1, 0).text = "c"
    table.cell(1, 1).text = "d"

    adapter._transfer_body_tables(cover, body)

    assert len(cover.tables) == 1
    new_table = cover.tables[0]
    assert new_table.cell(0, 0).text == "a"
    assert new_table.cell(1, 1).text == "d"


def _block_sequence(document):
    """Ordered list of block kinds for a document body: 'table' for each
    table, or the stripped text for each non-empty paragraph."""
    from docx.table import Table

    sequence = []
    for block in document.iter_inner_content():
        if isinstance(block, Table):
            sequence.append("table")
        else:
            text = block.text.strip()
            if text:
                sequence.append(text)
    return sequence


def test_body_tables_are_transferred_in_document_order(tmp_path):
    # A body with: Heading-1 "A", a table, paragraph "B". The table belongs
    # BETWEEN A and B, not appended after the last paragraph.
    body_docx = tmp_path / "body.docx"
    body = Document()
    body.add_paragraph("A", style="Heading 1")
    table = body.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "cell-x"
    table.cell(0, 1).text = "cell-y"
    body.add_paragraph("B")
    body.save(str(body_docx))

    adapter = PythonDocxAssemblyAdapter()
    config = {"structure": [{"type": "sections"}]}
    main = adapter._build_main_document(config, body_docx, None)

    sequence = _block_sequence(main)
    assert "A" in sequence and "B" in sequence and "table" in sequence
    a_idx = sequence.index("A")
    b_idx = sequence.index("B")
    table_idx = sequence.index("table")
    assert a_idx < table_idx < b_idx, sequence


def test_transferred_table_has_horizontal_only_borders_no_shading():
    from docs.infrastructure.docx.python_docx_audit_adapter import (
        table_has_vertical_borders_or_shading,
    )

    adapter = PythonDocxAssemblyAdapter()
    cover = Document()
    body = Document()
    table = body.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "head-a"
    table.cell(0, 1).text = "head-b"
    table.cell(1, 0).text = "line1\nline2"
    table.cell(1, 1).text = "d"

    adapter._transfer_body_tables(cover, body)

    new_table = cover.tables[0]
    xml = new_table._tbl.xml
    # Horizontal borders are declared.
    assert "<w:tblBorders" in xml
    borders = new_table._tbl.tblPr.find(qn("w:tblBorders"))
    assert borders is not None
    for edge in ("top", "bottom", "insideH"):
        el = borders.find(qn(f"w:{edge}"))
        assert el is not None and el.get(qn("w:val")) == "single"
    # Must PASS the institutional audit: no vertical borders, no shading.
    assert table_has_vertical_borders_or_shading(new_table) is False
    # Multi-line cell text is preserved.
    assert "line1" in new_table.cell(1, 0).text
    assert "line2" in new_table.cell(1, 0).text


def test_transfer_body_paragraphs_preserves_inline_images():
    adapter = PythonDocxAssemblyAdapter()
    cover = Document()
    body = Document()

    paragraph = body.add_paragraph()
    paragraph.add_run("before ")
    paragraph.add_run().add_picture(io.BytesIO(_PIXEL_PNG))
    paragraph.add_run(" after")

    adapter._transfer_body_paragraphs(cover, body, {}, {})

    # The transferred paragraph must carry the inline drawing, not just text.
    drawings = cover.element.body.findall(f".//{qn('w:drawing')}")
    assert len(drawings) >= 1
    # The image part must be embedded into the destination package.
    assert len(cover.inline_shapes) >= 1
    # Text runs around the image are still preserved in order.
    transferred_text = "".join(p.text for p in cover.paragraphs)
    assert "before" in transferred_text
    assert "after" in transferred_text
