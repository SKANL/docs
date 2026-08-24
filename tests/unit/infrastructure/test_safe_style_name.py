from __future__ import annotations

from typing import ClassVar

from docx import Document

from docs.infrastructure.docx.python_docx_assembly_adapter import safe_style_name


def test_safe_style_name_returns_preferred_when_already_available():
    document = Document()
    assert safe_style_name(document, "Heading 1") == "Heading 1"


def test_safe_style_name_maps_first_paragraph_to_no_spacing():
    document = Document()
    assert "First Paragraph" not in {s.name for s in document.styles}
    assert safe_style_name(document, "First Paragraph") == "No Spacing"


def test_safe_style_name_maps_compact_to_no_spacing():
    document = Document()
    assert safe_style_name(document, "Compact") == "No Spacing"


def test_safe_style_name_falls_back_to_normal_when_no_mapping_matches():
    document = Document()
    assert safe_style_name(document, "Some Unknown Style") == "Normal"


def test_safe_style_name_returns_none_when_neither_fallback_exists():
    class _FakeStyle:
        def __init__(self, name: str) -> None:
            self.name = name

    class _FakeDocument:
        styles: ClassVar[list] = [_FakeStyle("Custom Only")]

    assert safe_style_name(_FakeDocument(), "First Paragraph") is None


def test_safe_style_name_returns_none_for_none_preferred_style_without_fallback():
    class _FakeStyle:
        def __init__(self, name: str) -> None:
            self.name = name

    class _FakeDocument:
        styles: ClassVar[list] = [_FakeStyle("Custom Only")]

    assert safe_style_name(_FakeDocument(), None) is None


def test_safe_style_name_rejects_a_listed_style_the_document_cannot_address():
    # Found by the first CI run: 12 tests that pass on Windows failed on
    # Ubuntu with `KeyError: "no style with name 'Heading 1'"`, raised by
    # python-docx when the assembler applied a name `safe_style_name` had just
    # approved. Listing and addressing are not the same question -- iterating
    # `document.styles` can yield a name that `document.styles[name]` cannot
    # resolve, and which pandoc version produced the base document decides
    # whether that happens.
    #
    # `safe_style_name` exists precisely so a missing style degrades instead
    # of crashing. It was asking the wrong question.
    class _FakeStyle:
        def __init__(self, name: str) -> None:
            self.name = name

    class _FakeStyles:
        def __init__(self, listed: list[str], addressable: set[str]) -> None:
            self._listed = [_FakeStyle(n) for n in listed]
            self._addressable = addressable

        def __iter__(self):
            return iter(self._listed)

        def __getitem__(self, name: str):
            if name not in self._addressable:
                raise KeyError(f"no style with name '{name}'")
            return _FakeStyle(name)

    class _FakeDocument:
        def __init__(self, listed: list[str], addressable: set[str]) -> None:
            self.styles = _FakeStyles(listed, addressable)

    document = _FakeDocument(["Heading 1", "Normal"], addressable={"Normal"})

    assert safe_style_name(document, "Heading 1") == "Normal"


def test_safe_style_name_returns_none_when_nothing_is_addressable():
    class _FakeStyles:
        def __iter__(self):
            return iter([])

        def __getitem__(self, name: str):
            raise KeyError(name)

    class _FakeDocument:
        styles = _FakeStyles()

    assert safe_style_name(_FakeDocument(), "Heading 1") is None
