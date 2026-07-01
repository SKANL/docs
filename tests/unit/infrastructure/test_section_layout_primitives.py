from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches

from docs.infrastructure.docx.python_docx_assembly_adapter import (
    add_page_number_footer,
    apply_non_cover_section_layout,
    clear_story_part,
    set_section_page_number_start,
)


# --- apply_non_cover_section_layout -----------------------------------------


def test_apply_non_cover_section_layout_sets_letter_page_size():
    document = Document()
    section = document.sections[0]
    config = {"format": {"page_size": "letter"}}
    apply_non_cover_section_layout(section, config)
    assert section.page_width == Inches(8.5)
    assert section.page_height == Inches(11)


def test_apply_non_cover_section_layout_sets_configured_margins():
    document = Document()
    section = document.sections[0]
    config = {
        "format": {
            "page_margins_cm": {"non_cover": {"top": 2.5, "right": 3.0, "bottom": 2.5, "left": 3.0}},
        }
    }
    apply_non_cover_section_layout(section, config)
    # OOXML stores margins in twips (1/20 pt); centimeter values that aren't
    # exact multiples of a twip (~0.00176 cm) lose sub-twip precision on the
    # EMU -> twips -> EMU round trip performed by python-docx itself, so we
    # assert within one twip (635 EMU) rather than exact equality.
    one_twip_emu = 635
    assert abs(section.top_margin - Cm(2.5)) <= one_twip_emu
    assert abs(section.right_margin - Cm(3.0)) <= one_twip_emu
    assert abs(section.bottom_margin - Cm(2.5)) <= one_twip_emu
    assert abs(section.left_margin - Cm(3.0)) <= one_twip_emu


def test_apply_non_cover_section_layout_ignores_missing_margin_keys():
    document = Document()
    section = document.sections[0]
    original_left = section.left_margin
    apply_non_cover_section_layout(section, {"format": {"page_margins_cm": {"non_cover": {"top": 2.5}}}})
    assert section.left_margin == original_left


# --- clear_story_part --------------------------------------------------------


def test_clear_story_part_removes_existing_paragraphs_and_adds_one_empty():
    document = Document()
    section = document.sections[0]
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = "old content"
    section.footer.add_paragraph("more old content")
    clear_story_part(section.footer)
    assert len(section.footer.paragraphs) == 1
    assert section.footer.paragraphs[0].text == ""


# --- add_page_number_footer ---------------------------------------------------


def test_add_page_number_footer_sets_right_alignment_and_page_field(tmp_path):
    document = Document()
    section = document.sections[0]
    section.footer.is_linked_to_previous = False
    add_page_number_footer(section.footer)
    paragraph = section.footer.paragraphs[-1]
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    xml = paragraph._p.xml
    assert "PAGE" in xml
    assert 'w:fldCharType="begin"' in xml
    assert 'w:fldCharType="separate"' in xml
    assert 'w:fldCharType="end"' in xml


# --- set_section_page_number_start --------------------------------------------


def test_set_section_page_number_start_sets_start_and_format():
    document = Document()
    section = document.sections[0]
    set_section_page_number_start(section, 5, "lowerRoman")
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    assert pg_num_type.get(qn("w:start")) == "5"
    assert pg_num_type.get(qn("w:fmt")) == "lowerRoman"


def test_set_section_page_number_start_without_format_leaves_fmt_unset():
    document = Document()
    section = document.sections[0]
    set_section_page_number_start(section, 1)
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    assert pg_num_type.get(qn("w:start")) == "1"
    assert pg_num_type.get(qn("w:fmt")) is None


def test_set_section_page_number_start_reuses_existing_pg_num_type_element():
    document = Document()
    section = document.sections[0]
    set_section_page_number_start(section, 1, "decimal")
    set_section_page_number_start(section, 3, "lowerRoman")
    assert len(section._sectPr.findall(qn("w:pgNumType"))) == 1
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    assert pg_num_type.get(qn("w:start")) == "3"
    assert pg_num_type.get(qn("w:fmt")) == "lowerRoman"
